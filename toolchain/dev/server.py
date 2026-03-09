# SPDX-License-Identifier: MIT
# Copyright (c) Nico Ueberfeldt

import contextlib
import ast
import json
import os
import shutil
import subprocess
import sys
import re
import fnmatch
import uuid
import hashlib
import time
import math
import ssl
import urllib.error
import urllib.request
import urllib.parse
import html
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

try:
    import tomllib
except ModuleNotFoundError:  # pragma: no cover
    tomllib = None

try:
    import yaml
except ModuleNotFoundError:  # pragma: no cover
    yaml = None

try:
    import sympy as sp
except ModuleNotFoundError:  # pragma: no cover
    sp = None

try:
    import sqlparse
except ModuleNotFoundError:  # pragma: no cover
    sqlparse = None

try:
    from PIL import Image
except ModuleNotFoundError:  # pragma: no cover
    Image = None

try:
    import pytesseract
except ModuleNotFoundError:  # pragma: no cover
    pytesseract = None

try:
    from pypdf import PdfReader
except ModuleNotFoundError:  # pragma: no cover
    PdfReader = None

try:
    import docx
except ModuleNotFoundError:  # pragma: no cover
    docx = None

try:
    import openpyxl
except ModuleNotFoundError:  # pragma: no cover
    openpyxl = None

try:
    import xlrd
except ModuleNotFoundError:  # pragma: no cover
    xlrd = None

try:
    from tree_sitter_languages import get_parser as _ts_get_parser
except ModuleNotFoundError:  # pragma: no cover
    _ts_get_parser = None

from mcp.server.fastmcp import FastMCP
from starlette.applications import Starlette
from starlette.middleware.cors import CORSMiddleware
from starlette.responses import JSONResponse, PlainTextResponse
from starlette.routing import Mount, Route

REPO_PATH = Path(os.getenv("REPO_PATH", "/repo")).resolve()
HOST = os.getenv("HOST", "0.0.0.0")
PORT = int(os.getenv("PORT", "8000"))
MCP_TRANSPORT = os.getenv("MCP_TRANSPORT", "http").strip().lower()
ALLOW_MUTATIONS = os.getenv("ALLOW_MUTATIONS", "false").strip().lower() in {
    "1",
    "true",
    "yes",
    "on",
}
MAX_READ_BYTES = int(os.getenv("MAX_READ_BYTES", "262144"))
MAX_OUTPUT_CHARS = int(os.getenv("MAX_OUTPUT_CHARS", "200000"))
ALLOW_ORIGINS = [
    x.strip() for x in os.getenv("ALLOW_ORIGINS", "*").split(",") if x.strip()
]
LABS_DIR = Path("toolchain/dev/labs")
REPORTS_DIR = Path(".build/reports")
MEMORY_FILE = Path(".build/memory/context_memory.json")
FAILURE_MEMORY_FILE = Path(".build/memory/failure_memory.json")
TOKEN_BUDGET_FILE = Path(".build/memory/token_budget.json")
EDIT_TXN_DIR = Path(".build/transactions")
API_SNAPSHOT_FILE = Path(".build/reports/API_SURFACE.json")
REPO_INDEX_FILE = Path(".build/index/repo_index.json")
TOOL_CACHE_FILE = Path(".build/cache/tool_cache.json")
RESULT_STORE_FILE = Path(".build/cache/result_store.json")
OUTPUT_BASELINE_FILE = Path(".build/reports/TOOL_OUTPUT_BASELINE.json")
REUSE_SPDX_REPORT = Path(".build/reports/REUSE.spdx")
REUSE_LINT_REPORT = Path(".build/reports/REUSE_LINT.txt")
LOCAL_MODELS_DIR = Path(os.getenv("LOCAL_MODELS_DIR", "/models"))
LOCAL_EMBED_BACKEND = os.getenv("LOCAL_EMBED_BACKEND", "hash").strip().lower()
LOCAL_EMBED_MODEL = os.getenv("LOCAL_EMBED_MODEL", "").strip()
LOCAL_EMBED_DIM = int(os.getenv("LOCAL_EMBED_DIM", "256"))
LOCAL_INFER_BACKEND = os.getenv("LOCAL_INFER_BACKEND", "endpoint").strip().lower()
LOCAL_INFER_MODEL = os.getenv("LOCAL_INFER_MODEL", "").strip()
LOCAL_INFER_ENDPOINT = os.getenv(
    "LOCAL_INFER_ENDPOINT", "http://127.0.0.1:11434/api/generate"
).strip()
HOST_CA_CERT_FILE = os.getenv("HOST_CA_CERT_FILE", "").strip()
SAFE_COMMANDS = {"rg", "find", "sed", "awk", "jq", "git", "pytest", "reuse"}
SAFE_GIT_SUBCOMMANDS = {
    "status",
    "diff",
    "log",
    "show",
    "grep",
    "rev-parse",
    "branch",
    "ls-files",
}
OUTPUT_PROFILES = {"compact", "normal", "verbose"}

mcp = FastMCP(
    "git-repo-manager",
    instructions=(
        "Manage exactly one mounted Git repository and its files with minimal output. "
        "Prefer compact schemas, selective fields, pagination, and cached/indexed workflows. "
        "All paths are relative to the repository root."
    ),
)


def _trim_text(text: str, max_chars: int = MAX_OUTPUT_CHARS) -> str:
    if len(text) <= max_chars:
        return text
    return (
        text[:max_chars]
        + f"\n\n[truncated: output exceeded {max_chars} characters; original length={len(text)}]"
    )


def _ssl_context_for_url(url: str) -> ssl.SSLContext | None:
    if not url.lower().startswith("https://"):
        return None
    cafile = os.getenv("SSL_CERT_FILE", "").strip() or HOST_CA_CERT_FILE
    if cafile:
        p = Path(cafile)
        if p.is_file():
            return ssl.create_default_context(cafile=str(p))
    return ssl.create_default_context()


def _urlopen_with_host_certs(
    req: urllib.request.Request,
    timeout: int,
) -> Any:
    ctx = _ssl_context_for_url(req.full_url)
    if ctx is None:
        return urllib.request.urlopen(req, timeout=timeout)
    return urllib.request.urlopen(req, timeout=timeout, context=ctx)


def _html_to_text(raw_html: str) -> str:
    cleaned = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", raw_html)
    cleaned = re.sub(r"(?is)<[^>]+>", " ", cleaned)
    cleaned = html.unescape(cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def _ensure_repo_path_exists() -> None:
    REPO_PATH.mkdir(parents=True, exist_ok=True)


def _is_git_repo() -> bool:
    try:
        result = subprocess.run(
            ["git", "-C", str(REPO_PATH), "rev-parse", "--is-inside-work-tree"],
            check=False,
            capture_output=True,
            text=True,
        )
        return result.returncode == 0 and result.stdout.strip() == "true"
    except FileNotFoundError:
        raise RuntimeError("git executable not found inside container")


def _require_git_repo() -> None:
    _ensure_repo_path_exists()
    if not _is_git_repo():
        raise ValueError(f"{REPO_PATH} is not a Git working tree")


def _require_mutations() -> None:
    if not ALLOW_MUTATIONS:
        raise PermissionError(
            "Mutating operations are disabled. Set ALLOW_MUTATIONS=true to enable them."
        )


def _resolve_repo_path(rel_path: str = ".") -> Path:
    _ensure_repo_path_exists()
    candidate = (REPO_PATH / rel_path).resolve()
    try:
        candidate.relative_to(REPO_PATH)
    except ValueError as exc:
        raise ValueError("path escapes repository root") from exc
    return candidate


def _git(*args: str, check: bool = True) -> subprocess.CompletedProcess[str]:
    _ensure_repo_path_exists()
    try:
        result = subprocess.run(
            ["git", "-C", str(REPO_PATH), *args],
            check=False,
            capture_output=True,
            text=True,
        )
    except FileNotFoundError as exc:
        raise RuntimeError("git executable not found inside container") from exc

    if check and result.returncode != 0:
        stderr = result.stderr.strip()
        stdout = result.stdout.strip()
        msg = (
            stderr
            or stdout
            or f"git {' '.join(args)} failed with exit code {result.returncode}"
        )
        raise RuntimeError(msg)
    return result


def _normalize_paths(paths: list[str]) -> list[str]:
    normalized: list[str] = []
    for p in paths:
        resolved = _resolve_repo_path(p)
        normalized.append(str(resolved.relative_to(REPO_PATH)))
    return normalized


def _list_report_files(max_entries: int = 200) -> list[str]:
    reports_dir = _resolve_repo_path(str(REPORTS_DIR))
    if not reports_dir.exists():
        return []

    entries: list[str] = []
    for item in reports_dir.rglob("*"):
        if item.is_file():
            entries.append(str(item.relative_to(REPO_PATH)))
            if len(entries) >= max_entries:
                break
    entries.sort()
    return entries


def _is_hidden_rel_path(rel: Path) -> bool:
    return any(part.startswith(".") for part in rel.parts)


def _is_likely_binary(path: Path, max_file_bytes: int = 1048576) -> bool:
    try:
        size = path.stat().st_size
        if size > max_file_bytes:
            return True
        with path.open("rb") as f:
            chunk = f.read(8192)
        if b"\x00" in chunk:
            return True
    except OSError:
        return True
    return False


def _allowed_by_globs(
    rel_str: str,
    include_globs: list[str] | None = None,
    exclude_globs: list[str] | None = None,
) -> bool:
    if include_globs and not any(fnmatch.fnmatch(rel_str, g) for g in include_globs):
        return False
    if exclude_globs and any(fnmatch.fnmatch(rel_str, g) for g in exclude_globs):
        return False
    return True


def _iter_candidate_files(root: Path, recursive: bool) -> Any:
    if root.is_file():
        yield root
        return
    if recursive:
        for p in root.rglob("*"):
            if p.is_file():
                yield p
        return
    for p in root.glob("*"):
        if p.is_file():
            yield p


def _read_lines(path: Path, encoding: str = "utf-8") -> list[str]:
    return path.read_text(encoding=encoding, errors="replace").splitlines()


def _truncate_with_flag(text: str, max_chars: int) -> tuple[str, bool]:
    if max_chars < 1:
        raise ValueError("max_chars must be >= 1")
    if len(text) <= max_chars:
        return text, False
    marker = f"\n\n[truncated: exceeded {max_chars} chars]"
    keep = max(1, max_chars - len(marker))
    return text[:keep] + marker, True


def _read_pdf_text(path: Path, max_pages: int) -> tuple[str, dict[str, Any]]:
    if PdfReader is None:
        raise RuntimeError("pypdf is not installed")
    if max_pages < 1:
        raise ValueError("max_pages must be >= 1")
    reader = PdfReader(str(path))
    lines: list[str] = []
    page_count = len(reader.pages)
    for page in reader.pages[:max_pages]:
        extracted = page.extract_text() or ""
        if extracted:
            lines.append(extracted.strip())
    meta = {"page_count": page_count, "pages_read": min(page_count, max_pages)}
    return "\n\n".join(x for x in lines if x), meta


def _read_docx_text(path: Path) -> tuple[str, dict[str, Any]]:
    if docx is None:
        raise RuntimeError("python-docx is not installed")
    document = docx.Document(str(path))
    chunks: list[str] = []
    para_count = 0
    for p in document.paragraphs:
        t = p.text.strip()
        if t:
            chunks.append(t)
        para_count += 1
    table_cells = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    chunks.append(text)
                table_cells += 1
    return "\n".join(chunks), {"paragraph_count": para_count, "table_cells": table_cells}


def _read_doc_text(path: Path) -> tuple[str, dict[str, Any]]:
    antiword = shutil.which("antiword")
    if antiword:
        proc = subprocess.run(
            [antiword, str(path)],
            capture_output=True,
            text=True,
            check=False,
            encoding="utf-8",
            errors="replace",
        )
        output = (proc.stdout or "").strip()
        if output:
            return output, {"backend": "antiword", "exit_code": proc.returncode}
        return "", {"backend": "antiword", "exit_code": proc.returncode}

    # Fallback keeps behavior functional even without antiword.
    raw = path.read_bytes()
    text = raw.decode("latin-1", errors="replace")
    return text, {"backend": "latin1-fallback", "bytes_read": len(raw)}


def _read_xlsx_text(path: Path, max_rows_per_sheet: int) -> tuple[str, dict[str, Any]]:
    if openpyxl is None:
        raise RuntimeError("openpyxl is not installed")
    if max_rows_per_sheet < 1:
        raise ValueError("max_rows_per_sheet must be >= 1")
    wb = openpyxl.load_workbook(filename=str(path), read_only=True, data_only=True)
    chunks: list[str] = []
    total_rows = 0
    sheets_meta: list[dict[str, Any]] = []
    for ws in wb.worksheets:
        rows_read = 0
        for row in ws.iter_rows(values_only=True):
            if rows_read >= max_rows_per_sheet:
                break
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if values:
                chunks.append(" | ".join(values))
            rows_read += 1
        total_rows += rows_read
        sheets_meta.append({"name": ws.title, "rows_read": rows_read})
    return "\n".join(chunks), {"sheet_count": len(wb.worksheets), "rows_read": total_rows, "sheets": sheets_meta}


def _read_xls_text(path: Path, max_rows_per_sheet: int) -> tuple[str, dict[str, Any]]:
    if xlrd is None:
        raise RuntimeError("xlrd is not installed")
    if max_rows_per_sheet < 1:
        raise ValueError("max_rows_per_sheet must be >= 1")
    wb = xlrd.open_workbook(str(path))
    chunks: list[str] = []
    total_rows = 0
    sheets_meta: list[dict[str, Any]] = []
    for i in range(wb.nsheets):
        sheet = wb.sheet_by_index(i)
        rows_read = min(sheet.nrows, max_rows_per_sheet)
        for r in range(rows_read):
            row = sheet.row_values(r)
            values = [str(v).strip() for v in row if str(v).strip()]
            if values:
                chunks.append(" | ".join(values))
        total_rows += rows_read
        sheets_meta.append({"name": sheet.name, "rows_read": rows_read})
    return "\n".join(chunks), {"sheet_count": wb.nsheets, "rows_read": total_rows, "sheets": sheets_meta}


def _read_opendoc_text(path: Path, ext: str, max_rows_per_sheet: int) -> tuple[str, dict[str, Any]]:
    if max_rows_per_sheet < 1:
        raise ValueError("max_rows_per_sheet must be >= 1")
    try:
        with zipfile.ZipFile(path) as zf:
            raw = zf.read("content.xml")
    except KeyError as exc:
        raise RuntimeError("content.xml not found in OpenDocument file") from exc
    except zipfile.BadZipFile as exc:
        raise RuntimeError("invalid OpenDocument zip container") from exc

    try:
        root = ET.fromstring(raw)
    except ET.ParseError as exc:
        raise RuntimeError(f"invalid OpenDocument content.xml: {exc}") from exc

    ns = {
        "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
        "table": "urn:oasis:names:tc:opendocument:xmlns:table:1.0",
    }
    chunks: list[str] = []
    meta: dict[str, Any] = {"format": ext}

    if ext == ".ods":
        sheet_rows: list[dict[str, Any]] = []
        for table in root.findall(".//table:table", ns):
            table_name = table.attrib.get(f"{{{ns['table']}}}name", "")
            rows_read = 0
            for row in table.findall("table:table-row", ns):
                if rows_read >= max_rows_per_sheet:
                    break
                cells: list[str] = []
                for cell in row.findall("table:table-cell", ns):
                    ps = [p.text.strip() for p in cell.findall(".//text:p", ns) if p.text and p.text.strip()]
                    if ps:
                        cells.append(" ".join(ps))
                if cells:
                    chunks.append(" | ".join(cells))
                rows_read += 1
            sheet_rows.append({"name": table_name, "rows_read": rows_read})
        meta["sheet_count"] = len(sheet_rows)
        meta["sheets"] = sheet_rows
        meta["rows_read"] = sum(s["rows_read"] for s in sheet_rows)
        return "\n".join(chunks), meta

    # .odt / .odp text extraction
    paragraphs = [p.text.strip() for p in root.findall(".//text:p", ns) if p.text and p.text.strip()]
    chunks.extend(paragraphs)
    meta["paragraph_count"] = len(paragraphs)
    return "\n".join(chunks), meta


def _read_pptx_presentation(
    path: Path,
    max_slides: int,
    max_chars_per_slide: int,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    if max_slides < 1:
        raise ValueError("max_slides must be >= 1")
    if max_chars_per_slide < 1:
        raise ValueError("max_chars_per_slide must be >= 1")
    try:
        with zipfile.ZipFile(path) as zf:
            slide_files = sorted(
                [n for n in zf.namelist() if re.match(r"^ppt/slides/slide\d+\.xml$", n)],
                key=lambda n: int(re.search(r"slide(\d+)\.xml$", n).group(1)),  # type: ignore[union-attr]
            )
            slides: list[dict[str, Any]] = []
            for idx, slide_name in enumerate(slide_files[:max_slides], start=1):
                raw = zf.read(slide_name)
                root = ET.fromstring(raw)
                text_nodes = [t.text.strip() for t in root.findall(".//{*}t") if t.text and t.text.strip()]
                text = "\n".join(text_nodes)
                text, _ = _truncate_with_flag(text, max_chars=max_chars_per_slide)
                title = text_nodes[0] if text_nodes else f"Slide {idx}"
                slides.append(
                    {
                        "index": idx,
                        "title": title[:160],
                        "text": text,
                        "text_blocks": len(text_nodes),
                    }
                )
    except zipfile.BadZipFile as exc:
        raise RuntimeError("invalid .pptx container") from exc
    except ET.ParseError as exc:
        raise RuntimeError(f"invalid .pptx slide xml: {exc}") from exc
    return slides, {"slide_count": len(slide_files), "slides_read": len(slides)}


def _read_odp_presentation(
    path: Path,
    max_slides: int,
    max_chars_per_slide: int,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    if max_slides < 1:
        raise ValueError("max_slides must be >= 1")
    if max_chars_per_slide < 1:
        raise ValueError("max_chars_per_slide must be >= 1")
    try:
        with zipfile.ZipFile(path) as zf:
            raw = zf.read("content.xml")
    except KeyError as exc:
        raise RuntimeError("content.xml not found in .odp") from exc
    except zipfile.BadZipFile as exc:
        raise RuntimeError("invalid .odp container") from exc

    try:
        root = ET.fromstring(raw)
    except ET.ParseError as exc:
        raise RuntimeError(f"invalid .odp content xml: {exc}") from exc

    ns = {
        "draw": "urn:oasis:names:tc:opendocument:xmlns:drawing:1.0",
        "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
    }
    pages = root.findall(".//draw:page", ns)
    slides: list[dict[str, Any]] = []
    for idx, page in enumerate(pages[:max_slides], start=1):
        title = page.attrib.get(f"{{{ns['draw']}}}name", "") or f"Slide {idx}"
        lines = [p.text.strip() for p in page.findall(".//text:p", ns) if p.text and p.text.strip()]
        text = "\n".join(lines)
        text, _ = _truncate_with_flag(text, max_chars=max_chars_per_slide)
        slides.append(
            {
                "index": idx,
                "title": title[:160],
                "text": text,
                "text_blocks": len(lines),
            }
        )
    return slides, {"slide_count": len(pages), "slides_read": len(slides)}


def _read_ppt_legacy_text(path: Path, max_slides: int, max_chars_per_slide: int) -> tuple[list[dict[str, Any]], dict[str, Any], list[str]]:
    warnings: list[str] = []
    ppttxt = shutil.which("catppt") or shutil.which("ppttotext")
    text = ""
    if ppttxt:
        proc = subprocess.run(
            [ppttxt, str(path)],
            check=False,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
        text = (proc.stdout or "").strip()
    if not text:
        raw = path.read_bytes()
        ascii_chunks = re.findall(rb"[ -~]{4,}", raw)
        text = "\n".join(chunk.decode("latin-1", errors="replace") for chunk in ascii_chunks)
        warnings.append("Used lossy fallback extractor for .ppt (install catppt/ppttotext for better quality).")
    text, _ = _truncate_with_flag(text.strip(), max_chars=max(1, max_slides * max_chars_per_slide))
    slides = [{"index": 1, "title": "Slide 1", "text": text, "text_blocks": len(text.splitlines())}] if text else []
    return slides, {"slide_count": 1 if text else 0, "slides_read": len(slides)}, warnings


def _image_basic_features(path: Path) -> dict[str, Any]:
    features: dict[str, Any] = {
        "width": 0,
        "height": 0,
        "mode": "",
        "format": path.suffix.lower().lstrip("."),
        "aspect_ratio": 0.0,
        "mean_luma": None,
    }
    if Image is None:
        return features
    try:
        with Image.open(path) as img:
            width, height = img.size
            features["width"] = int(width)
            features["height"] = int(height)
            features["mode"] = str(img.mode)
            features["format"] = str((img.format or features["format"])).lower()
            features["aspect_ratio"] = round(width / height, 4) if height else 0.0
            gray = img.convert("L")
            stat = gray.resize((1, 1)).getpixel((0, 0))
            features["mean_luma"] = float(stat)
    except Exception:
        return features
    return features


def _node_display_name(node: ast.AST) -> str:
    if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)):
        return node.name
    if isinstance(node, ast.Call):
        return _ast_expr_name(node.func)
    if isinstance(node, ast.Import):
        return ", ".join(alias.name for alias in node.names)
    if isinstance(node, ast.ImportFrom):
        module = node.module or ""
        return f"{module}:{', '.join(alias.name for alias in node.names)}"
    return ""


def _ast_expr_name(node: ast.AST) -> str:
    if isinstance(node, ast.Name):
        return node.id
    if isinstance(node, ast.Attribute):
        left = _ast_expr_name(node.value)
        return f"{left}.{node.attr}" if left else node.attr
    if isinstance(node, ast.Call):
        return _ast_expr_name(node.func)
    return ""


def _guess_file_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".json":
        return "json"
    if suffix in {".toml"}:
        return "toml"
    if suffix in {".yaml", ".yml"}:
        return "yaml"
    raise ValueError("unsupported file type; expected .json, .toml, .yaml, or .yml")


def _parse_query_path(query: str) -> list[str | int]:
    query = query.strip()
    if not query:
        return []
    tokens: list[str | int] = []
    current = ""
    i = 0
    while i < len(query):
        ch = query[i]
        if ch == ".":
            if current:
                tokens.append(current)
                current = ""
            i += 1
            continue
        if ch == "[":
            if current:
                tokens.append(current)
                current = ""
            j = query.find("]", i)
            if j == -1:
                raise ValueError("invalid query path: missing closing ']'")
            raw_index = query[i + 1 : j].strip()
            if not raw_index.isdigit():
                raise ValueError("invalid query path: index must be a non-negative int")
            tokens.append(int(raw_index))
            i = j + 1
            continue
        current += ch
        i += 1
    if current:
        tokens.append(current)
    return tokens


def _query_value(data: Any, query: str) -> Any:
    value = data
    for token in _parse_query_path(query):
        if isinstance(token, int):
            if not isinstance(value, list):
                raise ValueError("query expected list while resolving index")
            if token < 0 or token >= len(value):
                raise ValueError("query index out of range")
            value = value[token]
            continue
        if not isinstance(value, dict):
            raise ValueError("query expected object while resolving key")
        if token not in value:
            raise ValueError(f"query key not found: {token}")
        value = value[token]
    return value


def _memory_load() -> dict[str, Any]:
    memory_path = _resolve_repo_path(str(MEMORY_FILE))
    if not memory_path.exists():
        return {"entries": []}
    try:
        payload = json.loads(memory_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {"entries": []}
    if not isinstance(payload, dict):
        return {"entries": []}
    entries = payload.get("entries", [])
    if not isinstance(entries, list):
        entries = []
    return {"entries": entries}


def _memory_save(payload: dict[str, Any]) -> None:
    memory_path = _resolve_repo_path(str(MEMORY_FILE))
    memory_path.parent.mkdir(parents=True, exist_ok=True)
    memory_path.write_text(
        json.dumps(payload, indent=2, sort_keys=True), encoding="utf-8"
    )


def _json_file_load(path: Path, default: Any) -> Any:
    file_path = _resolve_repo_path(str(path))
    if not file_path.exists():
        return default
    try:
        payload = json.loads(file_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return default
    return payload


def _json_file_save(path: Path, payload: Any) -> None:
    file_path = _resolve_repo_path(str(path))
    file_path.parent.mkdir(parents=True, exist_ok=True)
    file_path.write_text(json.dumps(payload, indent=2, sort_keys=True), encoding="utf-8")


def _token_budget_load() -> dict[str, Any]:
    payload = _json_file_load(
        TOKEN_BUDGET_FILE,
        {
            "max_output_chars": MAX_OUTPUT_CHARS,
            "default_output_profile": "compact",
        },
    )
    if not isinstance(payload, dict):
        return {"max_output_chars": MAX_OUTPUT_CHARS, "default_output_profile": "compact"}
    max_chars = payload.get("max_output_chars", MAX_OUTPUT_CHARS)
    profile = payload.get("default_output_profile", "compact")
    if not isinstance(max_chars, int) or max_chars < 1:
        max_chars = MAX_OUTPUT_CHARS
    if profile not in OUTPUT_PROFILES:
        profile = "compact"
    return {"max_output_chars": max_chars, "default_output_profile": profile}


def _token_budget_apply_max(max_chars: int | None) -> int:
    if isinstance(max_chars, int) and max_chars > 0:
        return max_chars
    return int(_token_budget_load()["max_output_chars"])


def _default_output_profile(output_profile: str | None) -> str:
    if output_profile and output_profile.strip():
        return _validate_output_profile(output_profile)
    return _validate_output_profile(_token_budget_load()["default_output_profile"])


def _paginate(items: list[Any], offset: int = 0, limit: int | None = None) -> list[Any]:
    if offset < 0:
        raise ValueError("offset must be >= 0")
    if limit is not None and limit < 1:
        raise ValueError("limit must be >= 1 when provided")
    if limit is None:
        return items[offset:]
    return items[offset : offset + limit]


def _select_fields(records: list[dict[str, Any]], fields: list[str] | None) -> list[dict[str, Any]]:
    if not fields:
        return records
    selected: list[dict[str, Any]] = []
    for row in records:
        selected.append({k: row.get(k) for k in fields if k in row})
    return selected


def _cache_load() -> dict[str, Any]:
    payload = _json_file_load(TOOL_CACHE_FILE, {"entries": {}})
    if not isinstance(payload, dict):
        return {"entries": {}}
    entries = payload.get("entries")
    if not isinstance(entries, dict):
        entries = {}
    return {"entries": entries}


def _cache_save(payload: dict[str, Any]) -> None:
    _json_file_save(TOOL_CACHE_FILE, payload)


def _cache_get(tool: str, key: str) -> Any | None:
    payload = _cache_load()
    tool_entries = payload["entries"].get(tool, {})
    if not isinstance(tool_entries, dict):
        return None
    row = tool_entries.get(key)
    if not isinstance(row, dict):
        return None
    return row.get("value")


def _cache_set(tool: str, key: str, value: Any, max_entries: int = 50) -> None:
    payload = _cache_load()
    entries = payload["entries"]
    tool_entries = entries.get(tool, {})
    if not isinstance(tool_entries, dict):
        tool_entries = {}
    tool_entries[key] = {"updated_at": _now_iso(), "value": value}
    if len(tool_entries) > max_entries:
        ordered = sorted(
            tool_entries.items(), key=lambda kv: str(kv[1].get("updated_at", "")), reverse=True
        )
        tool_entries = dict(ordered[:max_entries])
    entries[tool] = tool_entries
    payload["entries"] = entries
    _cache_save(payload)


def _cache_clear(tool: str | None = None) -> dict[str, Any]:
    payload = _cache_load()
    entries = payload.get("entries", {})
    if not isinstance(entries, dict):
        entries = {}
    removed = 0
    if tool:
        tool_entries = entries.pop(tool, {})
        if isinstance(tool_entries, dict):
            removed = len(tool_entries)
    else:
        for v in entries.values():
            if isinstance(v, dict):
                removed += len(v)
        entries = {}
    payload["entries"] = entries
    _cache_save(payload)
    return {"removed_entries": removed, "tool": tool or "*"}


def _cache_stats() -> dict[str, Any]:
    payload = _cache_load()
    entries = payload.get("entries", {})
    if not isinstance(entries, dict):
        entries = {}
    per_tool: dict[str, int] = {}
    total = 0
    for tool, rows in entries.items():
        if isinstance(rows, dict):
            per_tool[tool] = len(rows)
            total += len(rows)
    return {"total_entries": total, "tools": per_tool}


def _result_store_load() -> dict[str, Any]:
    payload = _json_file_load(RESULT_STORE_FILE, {"results": {}})
    if not isinstance(payload, dict):
        return {"results": {}}
    results = payload.get("results")
    if not isinstance(results, dict):
        results = {}
    return {"results": results}


def _result_store_save(payload: dict[str, Any]) -> None:
    _json_file_save(RESULT_STORE_FILE, payload)


def _result_store_put(tool: str, value: Any, max_entries: int = 200) -> str:
    payload = _result_store_load()
    results = payload["results"]
    rid = uuid.uuid4().hex[:16]
    results[rid] = {"tool": tool, "created_at": _now_iso(), "value": value}
    if len(results) > max_entries:
        ordered = sorted(
            results.items(), key=lambda kv: str(kv[1].get("created_at", "")), reverse=True
        )
        results = dict(ordered[:max_entries])
    payload["results"] = results
    _result_store_save(payload)
    return rid


def _result_store_get(result_id: str) -> dict[str, Any]:
    payload = _result_store_load()
    row = payload["results"].get(result_id)
    if not isinstance(row, dict):
        raise FileNotFoundError(f"result handle not found: {result_id}")
    return row


def _compress_table(records: list[dict[str, Any]]) -> dict[str, Any]:
    if not records:
        return {"columns": [], "rows": []}
    cols = sorted({k for r in records for k in r.keys()})
    rows = [[r.get(c) for c in cols] for r in records]
    return {"columns": cols, "rows": rows}


def _payload_size_bytes(value: Any) -> int:
    try:
        return len(json.dumps(value, ensure_ascii=True).encode("utf-8"))
    except Exception:
        return len(str(value).encode("utf-8"))


def _vec_l2(vec: list[float]) -> float:
    return math.sqrt(sum(x * x for x in vec))


def _vec_normalize(vec: list[float]) -> list[float]:
    n = _vec_l2(vec)
    if n == 0:
        return vec
    return [x / n for x in vec]


def _vec_cosine(a: list[float], b: list[float]) -> float:
    if not a or not b or len(a) != len(b):
        return 0.0
    denom = _vec_l2(a) * _vec_l2(b)
    if denom == 0:
        return 0.0
    return sum(x * y for x, y in zip(a, b)) / denom


def _hash_embed_one(text: str, dim: int = 256) -> list[float]:
    if dim < 8:
        raise ValueError("embedding dimension must be >= 8")
    vec = [0.0] * dim
    tokens = re.findall(r"[A-Za-z0-9_./:-]+", text.lower())
    if not tokens:
        tokens = [text.lower()]
    for tok in tokens:
        h = hashlib.sha256(tok.encode("utf-8")).digest()
        idx = int.from_bytes(h[:4], "big") % dim
        sign = -1.0 if (h[4] & 1) else 1.0
        vec[idx] += sign
    return _vec_normalize(vec)


def _local_embed_vectors(
    texts: list[str],
    backend: str = "auto",
    normalize: bool = True,
) -> tuple[str, list[list[float]]]:
    selected = backend.strip().lower()
    if selected == "auto":
        selected = LOCAL_EMBED_BACKEND or "hash"

    if selected in {"hash", "toy", "fallback"}:
        vectors = [_hash_embed_one(t, dim=LOCAL_EMBED_DIM) for t in texts]
        if normalize:
            vectors = [_vec_normalize(v) for v in vectors]
        return ("hash", vectors)

    if selected in {"sentence-transformers", "sentence_transformers"}:
        try:
            from sentence_transformers import SentenceTransformer  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise RuntimeError(
                "sentence-transformers backend requested but package is unavailable"
            ) from exc
        model_ref = LOCAL_EMBED_MODEL
        if not model_ref:
            raise RuntimeError("LOCAL_EMBED_MODEL is required for sentence-transformers backend")
        model = SentenceTransformer(model_ref, device="cpu")
        vectors_raw = model.encode(texts, normalize_embeddings=normalize)
        vectors = [[float(x) for x in row] for row in vectors_raw]
        return ("sentence-transformers", vectors)

    raise ValueError(
        "unsupported embed backend; expected one of: auto, hash, sentence-transformers"
    )


def _local_infer_via_endpoint(
    prompt: str,
    model: str,
    max_tokens: int,
    temperature: float,
    system: str = "",
) -> str:
    payload = {
        "model": model,
        "prompt": prompt,
        "system": system,
        "stream": False,
        "options": {"num_predict": max_tokens, "temperature": temperature},
    }
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        LOCAL_INFER_ENDPOINT,
        data=data,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with _urlopen_with_host_certs(req, timeout=60) as resp:
        body = resp.read().decode("utf-8", errors="replace")
    try:
        parsed = json.loads(body)
    except json.JSONDecodeError:
        return body
    if isinstance(parsed, dict):
        for key in ("response", "text", "output", "completion"):
            if isinstance(parsed.get(key), str):
                return parsed[key]
    return body


def _require_sympy() -> None:
    if sp is None:
        raise RuntimeError("sympy is not installed in this runtime")


def _math_expr(expr: str) -> Any:
    _require_sympy()
    return sp.sympify(expr)


def _math_steps_stub(mode: str, expr: str) -> list[str]:
    return [
        f"Parsed expression for mode '{mode}'.",
        "Applied symbolic transformation.",
        "Generated exact and numeric outputs.",
    ]


def _sql_normalize(query: str) -> str:
    if sqlparse is None:
        return " ".join(query.split())
    return sqlparse.format(query, keyword_case="upper", reindent=True)


def _extract_diff_lines(diff_text: str) -> list[str]:
    return [line for line in diff_text.splitlines() if line.startswith("+") or line.startswith("-")]


def _simple_translate(text: str, source_lang: str, target_lang: str) -> str:
    source = source_lang.lower()
    target = target_lang.lower()
    key = f"{source}->{target}"
    lexicons = {
        "en->de": {
            "hello": "hallo",
            "world": "welt",
            "error": "fehler",
            "success": "erfolg",
            "file": "datei",
            "test": "test",
        },
        "en->es": {
            "hello": "hola",
            "world": "mundo",
            "error": "error",
            "success": "exito",
            "file": "archivo",
            "test": "prueba",
        },
        "en->fr": {
            "hello": "bonjour",
            "world": "monde",
            "error": "erreur",
            "success": "succes",
            "file": "fichier",
            "test": "test",
        },
    }
    lex = lexicons.get(key)
    if not lex:
        return text
    tokens = re.findall(r"\w+|\W+", text)
    out: list[str] = []
    for t in tokens:
        low = t.lower()
        if low in lex:
            out.append(lex[low])
        else:
            out.append(t)
    return "".join(out)


def _diagram_fingerprint(paths: list[str]) -> str:
    h = hashlib.sha256()
    normalized = sorted(set(paths))
    for rel in normalized:
        p = _resolve_repo_path(rel)
        if not p.is_file():
            continue
        h.update(rel.encode("utf-8"))
        h.update(_file_sha256(p).encode("utf-8"))
    return h.hexdigest()


def _mermaid_sanitize_id(text: str) -> str:
    out = re.sub(r"[^A-Za-z0-9_]", "_", text)
    out = re.sub(r"_+", "_", out).strip("_")
    return out or "node"


def _adaptive_limit(requested: int, soft_cap: int = 500) -> int:
    if requested < 1:
        raise ValueError("requested limit must be >= 1")
    budget = _token_budget_load().get("max_output_chars", MAX_OUTPUT_CHARS)
    if not isinstance(budget, int):
        budget = MAX_OUTPUT_CHARS
    factor = 1.0
    if budget <= 100000:
        factor = 0.75
    if budget <= 50000:
        factor = 0.5
    if budget <= 25000:
        factor = 0.35
    cap = max(10, int(soft_cap * factor))
    return min(requested, cap)


def _fingerprint_path(
    root: Path,
    recursive: bool = True,
    suffixes: set[str] | None = None,
    max_files: int = 2000,
) -> str:
    hasher = hashlib.sha256()
    count = 0
    for p in _iter_candidate_files(root, recursive=recursive):
        if suffixes and p.suffix.lower() not in suffixes:
            continue
        rel = str(p.relative_to(REPO_PATH)).replace("\\", "/")
        hasher.update(rel.encode("utf-8"))
        try:
            hasher.update(_file_sha256(p).encode("utf-8"))
        except OSError:
            continue
        count += 1
        if count >= max_files:
            break
    return hasher.hexdigest()


def _failure_memory_load() -> dict[str, Any]:
    payload = _json_file_load(FAILURE_MEMORY_FILE, {"entries": []})
    if not isinstance(payload, dict):
        return {"entries": []}
    entries = payload.get("entries", [])
    if not isinstance(entries, list):
        entries = []
    return {"entries": entries}


def _failure_memory_save(payload: dict[str, Any]) -> None:
    _json_file_save(FAILURE_MEMORY_FILE, payload)


def _failure_record(
    command: list[str],
    stderr: str,
    stdout: str = "",
    category: str = "command",
    suggestion: str | None = None,
) -> None:
    payload = _failure_memory_load()
    entries = payload["entries"]
    entries.append(
        {
            "timestamp": _now_iso(),
            "command": command,
            "category": category,
            "stderr": _trim_text(stderr),
            "stdout": _trim_text(stdout),
            "suggestion": suggestion or "",
        }
    )
    payload["entries"] = entries[-500:]
    _failure_memory_save(payload)


def _tx_path(txn_id: str) -> Path:
    return _resolve_repo_path(str(EDIT_TXN_DIR / f"{txn_id}.json"))


def _tx_load(txn_id: str) -> dict[str, Any]:
    tx_path = _tx_path(txn_id)
    if not tx_path.is_file():
        raise FileNotFoundError(f"transaction not found: {txn_id}")
    try:
        payload = json.loads(tx_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise RuntimeError("transaction metadata is corrupted") from exc
    if not isinstance(payload, dict):
        raise RuntimeError("transaction metadata is invalid")
    return payload


def _tx_save(txn_id: str, payload: dict[str, Any]) -> None:
    tx_path = _tx_path(txn_id)
    tx_path.parent.mkdir(parents=True, exist_ok=True)
    tx_path.write_text(json.dumps(payload, indent=2, sort_keys=True), encoding="utf-8")


def _collect_python_symbols_top_level(
    source: str, rel_path: str, include_private: bool = False
) -> list[dict[str, Any]]:
    try:
        tree = ast.parse(source)
    except SyntaxError:
        return []
    symbols: list[dict[str, Any]] = []
    for node in tree.body:
        if not isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)):
            continue
        name = node.name
        if not include_private and name.startswith("_"):
            continue
        kind = "class"
        if isinstance(node, ast.FunctionDef):
            kind = "function"
        if isinstance(node, ast.AsyncFunctionDef):
            kind = "async_function"
        symbols.append(
            {
                "path": rel_path,
                "name": name,
                "kind": kind,
                "line_start": int(getattr(node, "lineno", 1)),
                "line_end": int(getattr(node, "end_lineno", getattr(node, "lineno", 1))),
            }
        )
    return symbols


def _file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _tree_sitter_language_for_ext(ext: str) -> str | None:
    mapping = {
        ".py": "python",
        ".js": "javascript",
        ".jsx": "javascript",
        ".ts": "typescript",
        ".tsx": "tsx",
        ".go": "go",
        ".rs": "rust",
    }
    return mapping.get(ext.lower())


def _tree_sitter_available() -> bool:
    return _ts_get_parser is not None


def _tree_sitter_parse_nodes(
    source: str,
    language: str,
    node_types: list[str] | None = None,
    max_nodes: int = 5000,
) -> list[dict[str, Any]]:
    if _ts_get_parser is None:
        raise RuntimeError("tree_sitter_languages is not installed")
    parser = _ts_get_parser(language)
    tree = parser.parse(source.encode("utf-8", errors="replace"))
    wanted = set(node_types or [])
    results: list[dict[str, Any]] = []

    stack = [tree.root_node]
    while stack and len(results) < max_nodes:
        node = stack.pop()
        if not wanted or node.type in wanted:
            results.append(
                {
                    "type": node.type,
                    "start_line": int(node.start_point[0]) + 1,
                    "start_column": int(node.start_point[1]) + 1,
                    "end_line": int(node.end_point[0]) + 1,
                    "end_column": int(node.end_point[1]) + 1,
                }
            )
        stack.extend(reversed(node.children))
    return results


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _to_iso_expiry(ttl_days: int | None) -> str | None:
    if ttl_days is None:
        return None
    if ttl_days < 1:
        raise ValueError("ttl_days must be >= 1")
    return (datetime.now(timezone.utc) + timedelta(days=ttl_days)).isoformat()


def _is_expired(expires_at: str | None, now: datetime) -> bool:
    if not expires_at:
        return False
    try:
        ts = datetime.fromisoformat(expires_at)
    except ValueError:
        return False
    if ts.tzinfo is None:
        ts = ts.replace(tzinfo=timezone.utc)
    return ts < now


def _collect_python_symbols(
    source: str, rel_path: str, include_private: bool
) -> list[dict[str, Any]]:
    try:
        tree = ast.parse(source)
    except SyntaxError:
        return []

    symbols: list[dict[str, Any]] = []
    for node in ast.walk(tree):
        if not isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)):
            continue
        name = node.name
        if not include_private and name.startswith("_"):
            continue
        kind = "class"
        if isinstance(node, ast.FunctionDef):
            kind = "function"
        if isinstance(node, ast.AsyncFunctionDef):
            kind = "async_function"
        symbols.append(
            {
                "path": rel_path,
                "name": name,
                "kind": kind,
                "line_start": int(getattr(node, "lineno", 1)),
                "line_end": int(getattr(node, "end_lineno", getattr(node, "lineno", 1))),
            }
        )
    return symbols


def _module_name_from_relpath(rel_path: Path) -> str:
    parts = list(rel_path.parts)
    if not parts:
        return ""
    if parts[-1] == "__init__.py":
        parts = parts[:-1]
    elif parts[-1].endswith(".py"):
        parts[-1] = parts[-1][:-3]
    return ".".join(parts)


def _import_candidates(module: str) -> list[str]:
    parts = module.split(".")
    candidates: list[str] = []
    for i in range(len(parts), 0, -1):
        prefix = ".".join(parts[:i])
        candidates.append(prefix)
    return candidates


def _validate_output_profile(output_profile: str) -> str:
    profile = output_profile.strip().lower()
    if profile not in OUTPUT_PROFILES:
        raise ValueError("output_profile must be one of: compact, normal, verbose")
    return profile


def _build_snippet(
    file_path: Path,
    start_line: int,
    end_line: int,
    context_before: int = 0,
    context_after: int = 0,
    encoding: str = "utf-8",
) -> dict[str, Any]:
    if start_line < 1 or end_line < 1:
        raise ValueError("start_line and end_line must be >= 1")
    if end_line < start_line:
        raise ValueError("end_line must be >= start_line")
    if context_before < 0 or context_after < 0:
        raise ValueError("context_before/context_after must be >= 0")
    if not file_path.is_file():
        raise FileNotFoundError(str(file_path.relative_to(REPO_PATH)))

    lines = _read_lines(file_path, encoding=encoding)
    total_lines = len(lines)
    from_line = max(1, start_line - context_before)
    to_line = min(total_lines, end_line + context_after)
    snippet_lines = lines[from_line - 1 : to_line]
    return {
        "path": str(file_path.relative_to(REPO_PATH)),
        "requested_start_line": start_line,
        "requested_end_line": end_line,
        "start_line": from_line,
        "end_line": to_line,
        "total_lines": total_lines,
        "content": "\n".join(snippet_lines),
    }


def _symbol_to_profile(symbol: dict[str, Any], profile: str) -> dict[str, Any]:
    if profile == "compact":
        return {
            "path": symbol["path"],
            "name": symbol["name"],
            "kind": symbol["kind"],
            "line_start": symbol["line_start"],
        }
    return symbol


def _match_to_profile(match: dict[str, Any], profile: str) -> dict[str, Any]:
    if profile == "compact":
        return {
            "path": match["path"],
            "line": match["line"],
            "column": match["column"],
            "match": match["match"],
        }
    return match


def _run_lab_script(script_name: str, args: list[str]) -> dict[str, Any]:
    _require_mutations()
    _require_git_repo()

    script_rel = str(LABS_DIR / script_name)
    script_path = _resolve_repo_path(script_rel)
    if not script_path.is_file():
        raise FileNotFoundError(script_rel)

    proc = subprocess.run(
        [sys.executable, str(script_path), *args],
        check=False,
        capture_output=True,
        text=True,
        cwd=str(REPO_PATH),
    )

    stdout = _trim_text(proc.stdout.strip())
    stderr = _trim_text(proc.stderr.strip())
    result: dict[str, Any] = {
        "script": script_rel,
        "args": args,
        "exit_code": proc.returncode,
        "ok": proc.returncode == 0,
        "stdout": stdout,
        "stderr": stderr,
        "reports": _list_report_files(),
    }

    if proc.returncode != 0:
        msg = (
            stderr or stdout or f"{script_name} failed with exit code {proc.returncode}"
        )
        raise RuntimeError(msg)

    return result


def _chunk_strings(values: list[str], chunk_size: int) -> list[list[str]]:
    if chunk_size < 1:
        raise ValueError("chunk_size must be >= 1")
    return [values[i : i + chunk_size] for i in range(0, len(values), chunk_size)]


def _require_reuse_cli() -> None:
    if shutil.which("reuse") is None:
        raise RuntimeError("reuse CLI not found; install python package 'reuse'")


def _run_reuse(args: list[str], timeout_seconds: int = 120) -> dict[str, Any]:
    _require_reuse_cli()
    if timeout_seconds < 1:
        raise ValueError("timeout_seconds must be >= 1")
    proc = subprocess.run(
        ["reuse", *args],
        cwd=str(REPO_PATH),
        check=False,
        capture_output=True,
        text=True,
        timeout=timeout_seconds,
    )
    return {
        "ok": proc.returncode == 0,
        "exit_code": proc.returncode,
        "command": ["reuse", *args],
        "stdout": _trim_text(proc.stdout.strip()),
        "stderr": _trim_text(proc.stderr.strip()),
    }


def _collect_spdx_license_ids(path: str = ".", recursive: bool = True) -> list[str]:
    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)
    found: set[str] = set()
    matcher = re.compile(r"SPDX-License-Identifier:\s*(.+)")
    token_re = re.compile(r"[A-Za-z0-9.\-+]+")
    keywords = {"AND", "OR", "WITH"}
    for candidate in _iter_candidate_files(root, recursive=recursive):
        rel = candidate.relative_to(REPO_PATH)
        rel_str = str(rel).replace("\\", "/")
        if rel_str.startswith(".git/") or rel_str.startswith(".build/"):
            continue
        if _is_likely_binary(candidate):
            continue
        text = candidate.read_text(encoding="utf-8", errors="replace")
        for line in text.splitlines()[:120]:
            match = matcher.search(line)
            if not match:
                continue
            expr = match.group(1)
            for token in token_re.findall(expr):
                if token in keywords or token.startswith("LicenseRef-"):
                    continue
                found.add(token)
    return sorted(found)


def _collect_missing_spdx_headers(
    path: str = ".",
    recursive: bool = True,
    include_globs: list[str] | None = None,
    exclude_globs: list[str] | None = None,
    max_files: int = 5000,
) -> list[str]:
    if max_files < 1:
        raise ValueError("max_files must be >= 1")
    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)
    missing: list[str] = []
    for candidate in _iter_candidate_files(root, recursive=recursive):
        rel = candidate.relative_to(REPO_PATH)
        rel_str = str(rel).replace("\\", "/")
        if rel_str.startswith(".git/") or rel_str.startswith(".build/") or rel_str.startswith("LICENSES/"):
            continue
        if not _allowed_by_globs(rel_str, include_globs=include_globs, exclude_globs=exclude_globs):
            continue
        if _is_likely_binary(candidate):
            continue
        try:
            lines = _read_lines(candidate, encoding="utf-8")
        except OSError:
            continue
        window = "\n".join(lines[:40])
        if "SPDX-License-Identifier:" not in window:
            missing.append(rel_str)
            if len(missing) >= max_files:
                break
    return missing


@mcp.tool()
def repo_info() -> dict[str, Any]:
    """Return repository state and server settings."""
    _ensure_repo_path_exists()

    info: dict[str, Any] = {
        "repo_path": str(REPO_PATH),
        "repo_exists": REPO_PATH.exists(),
        "is_git_repo": _is_git_repo(),
        "allow_mutations": ALLOW_MUTATIONS,
        "transport": MCP_TRANSPORT,
        "max_read_bytes": MAX_READ_BYTES,
        "max_output_chars": MAX_OUTPUT_CHARS,
    }

    if info["is_git_repo"]:
        info["current_branch"] = _git("branch", "--show-current").stdout.strip()
        info["head"] = _git("rev-parse", "HEAD").stdout.strip()
        status = _git("status", "--porcelain").stdout.strip()
        info["dirty"] = bool(status)

    return info


@mcp.tool()
def git_init(initial_branch: str = "main") -> dict[str, str]:
    """Initialize a Git repository in the mounted directory."""
    _require_mutations()
    _ensure_repo_path_exists()
    if _is_git_repo():
        raise ValueError(f"{REPO_PATH} is already a Git working tree")

    _git("init", "-b", initial_branch)
    return {
        "repo_path": str(REPO_PATH),
        "message": f"initialized repository with initial branch '{initial_branch}'",
    }


@mcp.tool()
def list_files(
    path: str = ".",
    recursive: bool = True,
    include_hidden: bool = False,
    max_entries: int = 1000,
) -> list[str]:
    """List files and directories under a repository-relative path."""
    if max_entries < 1:
        raise ValueError("max_entries must be >= 1")

    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)

    entries: list[str] = []

    def include_item(p: Path) -> bool:
        rel = p.relative_to(REPO_PATH)
        if include_hidden:
            return True
        return not any(part.startswith(".") for part in rel.parts)

    if root.is_file():
        if include_item(root):
            return [str(root.relative_to(REPO_PATH))]
        return []

    iterator = root.rglob("*") if recursive else root.glob("*")
    for item in iterator:
        if not include_item(item):
            continue
        rel = str(item.relative_to(REPO_PATH))
        if item.is_dir():
            rel += "/"
        entries.append(rel)
        if len(entries) >= max_entries:
            break

    entries.sort()
    return entries


@mcp.tool()
def read_file(
    path: str, encoding: str = "utf-8", max_bytes: int = MAX_READ_BYTES
) -> str:
    """Read a UTF-8 text file from the repository."""
    if max_bytes < 1:
        raise ValueError("max_bytes must be >= 1")

    file_path = _resolve_repo_path(path)
    if not file_path.is_file():
        raise FileNotFoundError(path)

    size = file_path.stat().st_size
    if size > max_bytes:
        raise ValueError(f"file is too large ({size} bytes > {max_bytes} bytes)")

    data = file_path.read_text(encoding=encoding)
    return _trim_text(data)


@mcp.tool()
def read_document(
    path: str,
    max_chars: int = 20000,
    max_pages: int = 20,
    max_rows_per_sheet: int = 200,
    output_profile: str | None = None,
) -> dict[str, Any]:
    """Read document formats: .pdf, .doc, .docx, .xls, .xlsx, .odt, .ods, .odp."""
    profile = _default_output_profile(output_profile)
    file_path = _resolve_repo_path(path)
    if not file_path.is_file():
        raise FileNotFoundError(path)
    ext = file_path.suffix.lower()
    if ext not in {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".odt", ".ods", ".odp"}:
        raise ValueError(
            "unsupported extension; expected .pdf, .doc, .docx, .xls, .xlsx, .odt, .ods, or .odp"
        )

    warnings: list[str] = []
    metadata: dict[str, Any] = {"extension": ext}
    text = ""

    try:
        if ext == ".pdf":
            text, extra = _read_pdf_text(file_path, max_pages=max_pages)
        elif ext == ".docx":
            text, extra = _read_docx_text(file_path)
        elif ext == ".doc":
            text, extra = _read_doc_text(file_path)
            if extra.get("backend") != "antiword":
                warnings.append("antiword not available; used lossy latin-1 fallback for .doc")
        elif ext == ".xlsx":
            text, extra = _read_xlsx_text(file_path, max_rows_per_sheet=max_rows_per_sheet)
        elif ext in {".odt", ".ods", ".odp"}:
            text, extra = _read_opendoc_text(
                file_path,
                ext=ext,
                max_rows_per_sheet=max_rows_per_sheet,
            )
        else:
            text, extra = _read_xls_text(file_path, max_rows_per_sheet=max_rows_per_sheet)
    except Exception as exc:
        raise RuntimeError(f"failed to parse {ext} document: {exc}") from exc

    metadata.update(extra)
    text = text.strip()
    text, truncated = _truncate_with_flag(text, max_chars=max_chars)

    result = {
        "schema": "read_document.v1",
        "path": str(file_path.relative_to(REPO_PATH)),
        "extension": ext,
        "chars": len(text),
        "truncated": truncated,
        "warnings": warnings,
        "metadata": metadata if profile != "compact" else {},
        "text": text,
    }
    if profile == "compact":
        return {
            "schema": "read_document.compact.v1",
            "path": result["path"],
            "extension": ext,
            "chars": result["chars"],
            "truncated": truncated,
            "warnings": warnings,
            "text": text,
        }
    return result


@mcp.tool()
def browse_web(
    url: str,
    timeout_seconds: int = 15,
    max_bytes: int = 300000,
    max_chars: int = 12000,
    extract_text: bool = True,
    output_profile: str | None = None,
) -> dict[str, Any]:
    """Fetch a web page/document over HTTP(S) using host/system certificate store."""
    if timeout_seconds < 1:
        raise ValueError("timeout_seconds must be >= 1")
    if max_bytes < 1:
        raise ValueError("max_bytes must be >= 1")
    if max_chars < 1:
        raise ValueError("max_chars must be >= 1")
    parsed = urllib.parse.urlparse(url)
    if parsed.scheme.lower() not in {"http", "https"}:
        raise ValueError("url scheme must be http or https")

    profile = _default_output_profile(output_profile)
    req = urllib.request.Request(
        url,
        method="GET",
        headers={
            "User-Agent": "repo-git-mcp/1.0 (+browse_web)",
            "Accept": "text/html,application/xhtml+xml,application/xml,text/plain;q=0.9,*/*;q=0.8",
        },
    )
    try:
        with _urlopen_with_host_certs(req, timeout=timeout_seconds) as resp:
            status = int(getattr(resp, "status", 200))
            final_url = str(getattr(resp, "url", url))
            content_type = str(resp.headers.get("Content-Type", ""))
            raw = resp.read(max_bytes + 1)
    except urllib.error.URLError as exc:
        raise RuntimeError(f"browse failed: {exc}") from exc

    truncated_bytes = len(raw) > max_bytes
    if truncated_bytes:
        raw = raw[:max_bytes]
    charset_match = re.search(r"charset=([A-Za-z0-9._-]+)", content_type, flags=re.IGNORECASE)
    encoding = charset_match.group(1) if charset_match else "utf-8"
    text_raw = raw.decode(encoding, errors="replace")
    text = _html_to_text(text_raw) if extract_text else text_raw
    text, truncated_chars = _truncate_with_flag(text, max_chars=max_chars)

    result = {
        "schema": "browse_web.v1",
        "url": url,
        "final_url": final_url,
        "status": status,
        "content_type": content_type,
        "encoding": encoding,
        "extract_text": extract_text,
        "bytes_read": len(raw),
        "truncated_bytes": truncated_bytes,
        "truncated_chars": truncated_chars,
        "text": text,
    }
    if profile == "compact":
        return {
            "schema": "browse_web.compact.v1",
            "url": result["url"],
            "final_url": result["final_url"],
            "status": result["status"],
            "content_type": result["content_type"],
            "truncated": bool(truncated_bytes or truncated_chars),
            "text": result["text"],
        }
    return result


@mcp.tool()
def interpret_presentation(
    path: str,
    max_slides: int = 50,
    max_chars_per_slide: int = 1200,
    use_local_model: bool = True,
    output_profile: str | None = None,
) -> dict[str, Any]:
    """Interpret presentation files (.pptx, .ppt, .odp) and optionally summarize with a local model."""
    profile = _default_output_profile(output_profile)
    file_path = _resolve_repo_path(path)
    if not file_path.is_file():
        raise FileNotFoundError(path)
    ext = file_path.suffix.lower()
    if ext not in {".pptx", ".ppt", ".odp"}:
        raise ValueError("unsupported extension; expected .pptx, .ppt, or .odp")

    warnings: list[str] = []
    if ext == ".pptx":
        slides, meta = _read_pptx_presentation(
            file_path, max_slides=max_slides, max_chars_per_slide=max_chars_per_slide
        )
    elif ext == ".odp":
        slides, meta = _read_odp_presentation(
            file_path, max_slides=max_slides, max_chars_per_slide=max_chars_per_slide
        )
    else:
        slides, meta, legacy_warnings = _read_ppt_legacy_text(
            file_path, max_slides=max_slides, max_chars_per_slide=max_chars_per_slide
        )
        warnings.extend(legacy_warnings)

    interpreted_summary = ""
    model_used = ""
    if use_local_model and slides:
        joined = "\n\n".join(
            f"Slide {s['index']} - {s['title']}\n{s['text']}" for s in slides[:12]
        )
        prompt = (
            "Summarize this presentation in <=8 bullets. "
            "Include objective, key points, risks, and action items.\n\n"
            f"{joined}"
        )
        try:
            inferred = local_infer(
                prompt=prompt,
                task="presentation_summary",
                backend="auto",
                output_profile="compact",
                max_tokens=400,
            )
            interpreted_summary = str(inferred.get("output", "")).strip()
            model_used = str(inferred.get("backend", ""))
        except Exception as exc:
            warnings.append(f"local model summary failed: {exc}")

    if not interpreted_summary:
        titles = [str(s.get("title", "")).strip() for s in slides if str(s.get("title", "")).strip()]
        if titles:
            interpreted_summary = "Slides cover: " + ", ".join(titles[:8])
            if len(titles) > 8:
                interpreted_summary += ", ..."
        else:
            interpreted_summary = "No textual content extracted from presentation."

    result = {
        "schema": "interpret_presentation.v1",
        "path": str(file_path.relative_to(REPO_PATH)),
        "extension": ext,
        "slide_count": int(meta.get("slide_count", len(slides))),
        "slides_read": int(meta.get("slides_read", len(slides))),
        "used_local_model": bool(model_used),
        "model_backend": model_used,
        "warnings": warnings,
        "summary": interpreted_summary,
        "slides": slides,
    }
    if profile == "compact":
        return {
            "schema": "interpret_presentation.compact.v1",
            "path": result["path"],
            "extension": ext,
            "slide_count": result["slide_count"],
            "slides_read": result["slides_read"],
            "used_local_model": result["used_local_model"],
            "warnings": warnings,
            "summary": interpreted_summary,
            "slides": [{"index": s["index"], "title": s["title"]} for s in slides[:50]],
        }
    return result


@mcp.tool()
def write_file(
    path: str,
    content: str,
    overwrite: bool = True,
    create_dirs: bool = True,
    encoding: str = "utf-8",
) -> dict[str, Any]:
    """Write a text file into the repository."""
    _require_mutations()
    file_path = _resolve_repo_path(path)

    if file_path.exists() and file_path.is_dir():
        raise IsADirectoryError(path)
    if file_path.exists() and not overwrite:
        raise FileExistsError(path)

    existed_before = file_path.exists()

    if create_dirs:
        file_path.parent.mkdir(parents=True, exist_ok=True)

    file_path.write_text(content, encoding=encoding)
    return {
        "path": str(file_path.relative_to(REPO_PATH)),
        "bytes_written": len(content.encode(encoding)),
        "existed_before": existed_before,
    }


@mcp.tool()
def delete_path(path: str, recursive: bool = False) -> dict[str, str]:
    """Delete a file or directory inside the repository."""
    _require_mutations()
    target = _resolve_repo_path(path)
    if not target.exists():
        raise FileNotFoundError(path)

    rel = str(target.relative_to(REPO_PATH))
    if target.is_dir():
        if recursive:
            shutil.rmtree(target)
        else:
            target.rmdir()
    else:
        target.unlink()
    return {"deleted": rel}


@mcp.tool()
def move_path(
    source: str,
    destination: str,
    overwrite: bool = False,
    create_dirs: bool = True,
) -> dict[str, str]:
    """Move or rename a file or directory inside the repository."""
    _require_mutations()
    src = _resolve_repo_path(source)
    dst = _resolve_repo_path(destination)

    if not src.exists():
        raise FileNotFoundError(source)
    if dst.exists() and not overwrite:
        raise FileExistsError(destination)
    if create_dirs:
        dst.parent.mkdir(parents=True, exist_ok=True)
    if dst.exists() and overwrite:
        if dst.is_dir():
            shutil.rmtree(dst)
        else:
            dst.unlink()

    shutil.move(str(src), str(dst))
    return {
        "source": str(src.relative_to(REPO_PATH)),
        "destination": str(dst.relative_to(REPO_PATH)),
    }


@mcp.tool()
def git_status(short: bool = True) -> str:
    """Return git status."""
    _require_git_repo()
    args = ["status"]
    if short:
        args.append("--short")
    return _trim_text(_git(*args).stdout)


@mcp.tool()
def git_diff(
    ref: str | None = None,
    pathspec: str | None = None,
    staged: bool = False,
) -> str:
    """Return a git diff against the working tree, index, or a ref."""
    _require_git_repo()
    args = ["diff"]
    if staged:
        args.append("--staged")
    if ref:
        args.append(ref)
    if pathspec:
        _resolve_repo_path(pathspec)
        args.extend(["--", pathspec])
    return _trim_text(_git(*args).stdout)


@mcp.tool()
def git_log(limit: int = 20, ref: str = "HEAD") -> str:
    """Return recent commit history."""
    _require_git_repo()
    if limit < 1:
        raise ValueError("limit must be >= 1")
    return _trim_text(
        _git(
            "log",
            f"--max-count={limit}",
            "--decorate",
            "--graph",
            "--oneline",
            ref,
        ).stdout
    )


@mcp.tool()
def git_show(ref: str = "HEAD", path: str | None = None) -> str:
    """Show a commit, object, or file content from Git history."""
    _require_git_repo()
    target = ref
    if path:
        _resolve_repo_path(path)
        target = f"{ref}:{path}"
    return _trim_text(_git("show", target).stdout)


@mcp.tool()
def git_add(paths: list[str]) -> dict[str, Any]:
    """Stage one or more repository-relative paths."""
    _require_mutations()
    _require_git_repo()
    if not paths:
        raise ValueError("paths must not be empty")
    normalized = _normalize_paths(paths)
    _git("add", "--", *normalized)
    return {"staged": normalized}


@mcp.tool()
def git_restore(paths: list[str], staged: bool = False) -> dict[str, Any]:
    """Restore paths from HEAD or unstage them."""
    _require_mutations()
    _require_git_repo()
    if not paths:
        raise ValueError("paths must not be empty")
    normalized = _normalize_paths(paths)
    args = ["restore"]
    if staged:
        args.append("--staged")
    args.extend(["--", *normalized])
    _git(*args)
    return {"restored": normalized, "staged": staged}


@mcp.tool()
def git_commit(message: str, allow_empty: bool = False) -> dict[str, str]:
    """Create a commit from staged changes."""
    _require_mutations()
    _require_git_repo()
    if not message.strip():
        raise ValueError("message must not be empty")

    args = ["commit", "-m", message]
    if allow_empty:
        args.append("--allow-empty")
    _git(*args)
    return {
        "commit": _git("rev-parse", "HEAD").stdout.strip(),
        "summary": _git("log", "-1", "--oneline").stdout.strip(),
    }


@mcp.tool()
def git_checkout(ref: str, create_branch: bool = False) -> dict[str, str]:
    """Checkout an existing ref or create a new branch from the current HEAD."""
    _require_mutations()
    _require_git_repo()
    if not ref.strip():
        raise ValueError("ref must not be empty")

    if create_branch:
        _git("checkout", "-b", ref)
    else:
        _git("checkout", ref)
    return {"current_branch": _git("branch", "--show-current").stdout.strip()}


@mcp.tool()
def git_create_branch(name: str, checkout: bool = True) -> dict[str, str]:
    """Create a branch, optionally switching to it immediately."""
    _require_mutations()
    _require_git_repo()
    if not name.strip():
        raise ValueError("name must not be empty")

    if checkout:
        _git("checkout", "-b", name)
    else:
        _git("branch", name)
    return {
        "branch": name,
        "checked_out": str(checkout).lower(),
        "current_branch": _git("branch", "--show-current").stdout.strip(),
    }


@mcp.tool()
def git_fetch(remote: str = "origin", prune: bool = False) -> str:
    """Fetch from a remote."""
    _require_mutations()
    _require_git_repo()
    args = ["fetch"]
    if prune:
        args.append("--prune")
    args.append(remote)
    result = _git(*args)
    return _trim_text(result.stdout + result.stderr)


@mcp.tool()
def git_pull(
    remote: str = "origin", branch: str | None = None, rebase: bool = False
) -> str:
    """Pull from a remote."""
    _require_mutations()
    _require_git_repo()
    args = ["pull"]
    if rebase:
        args.append("--rebase")
    args.append(remote)
    if branch:
        args.append(branch)
    result = _git(*args)
    return _trim_text(result.stdout + result.stderr)


@mcp.tool()
def git_push(
    remote: str = "origin",
    branch: str | None = None,
    set_upstream: bool = False,
) -> str:
    """Push to a remote."""
    _require_mutations()
    _require_git_repo()
    args = ["push"]
    if set_upstream:
        args.append("-u")
    args.append(remote)
    if branch:
        args.append(branch)
    result = _git(*args)
    return _trim_text(result.stdout + result.stderr)


@mcp.tool()
def lab_release_rehearsal(
    config_path: str = ".config/dev/labs/release_rehearsal.json",
    allow_dirty: bool = False,
    keep_branch: bool = False,
) -> dict[str, Any]:
    """Run release rehearsal lab and write report(s) under .build/reports."""
    _resolve_repo_path(config_path)
    args = ["--config", config_path]
    if allow_dirty:
        args.append("--allow-dirty")
    if keep_branch:
        args.append("--keep-branch")
    return _run_lab_script("release_rehearsal.py", args)


@mcp.tool()
def lab_refactor_tournament(
    config_path: str = ".config/dev/labs/refactor_tournament.json",
    allow_dirty: bool = False,
    keep_branches: bool = False,
) -> dict[str, Any]:
    """Run refactor tournament lab and write report(s) under .build/reports."""
    _resolve_repo_path(config_path)
    args = ["--config", config_path]
    if allow_dirty:
        args.append("--allow-dirty")
    if keep_branches:
        args.append("--keep-branches")
    return _run_lab_script("refactor_tournament.py", args)


@mcp.tool()
def lab_policy_gatekeeper(
    config_path: str = ".config/dev/labs/policy_gatekeeper.json",
    changed_ref: str = "HEAD",
    report_path: str = ".build/reports/POLICY_GATEKEEPER.md",
) -> dict[str, Any]:
    """Run policy-as-code gatekeeper checks."""
    _resolve_repo_path(config_path)
    _resolve_repo_path(report_path)
    args = [
        "--config",
        config_path,
        "--changed-ref",
        changed_ref,
        "--report-path",
        report_path,
    ]
    return _run_lab_script("policy_gatekeeper.py", args)


@mcp.tool()
def lab_branch_swarm(
    config_path: str = ".config/dev/labs/branch_swarm_lab.json",
    allow_dirty: bool = False,
    keep_branches: bool = False,
) -> dict[str, Any]:
    """Run branch swarm benchmark lab."""
    _resolve_repo_path(config_path)
    args = ["--config", config_path]
    if allow_dirty:
        args.append("--allow-dirty")
    if keep_branches:
        args.append("--keep-branches")
    return _run_lab_script("branch_swarm_lab.py", args)


@mcp.tool()
def lab_narrated_pr(
    base: str = "HEAD~1",
    head: str = "HEAD",
    output_path: str = ".build/reports/PR_PACKET.md",
) -> dict[str, Any]:
    """Generate a narrated PR packet for a commit range."""
    _resolve_repo_path(output_path)
    args = ["--base", base, "--head", head, "--output", output_path]
    return _run_lab_script("narrated_pr_generator.py", args)


@mcp.tool()
def lab_repo_digital_twin(
    json_path: str = ".build/reports/REPO_DIGITAL_TWIN.json",
    markdown_path: str = ".build/reports/REPO_DIGITAL_TWIN.md",
    max_files: int = 1000,
    hotspot_limit: int = 20,
) -> dict[str, Any]:
    """Generate repo digital twin JSON + markdown snapshots."""
    if max_files < 1:
        raise ValueError("max_files must be >= 1")
    if hotspot_limit < 1:
        raise ValueError("hotspot_limit must be >= 1")
    _resolve_repo_path(json_path)
    _resolve_repo_path(markdown_path)
    args = [
        "--json",
        json_path,
        "--md",
        markdown_path,
        "--max-files",
        str(max_files),
        "--hotspot-limit",
        str(hotspot_limit),
    ]
    return _run_lab_script("repo_digital_twin.py", args)


@mcp.tool()
def license_monitor(
    path: str = ".",
    recursive: bool = True,
    include_globs: list[str] | None = None,
    exclude_globs: list[str] | None = None,
    run_reuse_lint: bool = True,
    generate_spdx: bool = True,
    spdx_output_path: str = str(REUSE_SPDX_REPORT),
    lint_report_path: str = str(REUSE_LINT_REPORT),
    auto_fix_headers: bool = False,
    default_license: str = "MIT",
    copyright_owner: str = "Project Contributors",
    download_missing_licenses: bool = False,
    max_missing_files: int = 5000,
) -> dict[str, Any]:
    """Check REUSE/FOSS license compliance and optionally auto-fix missing metadata."""
    if max_missing_files < 1:
        raise ValueError("max_missing_files must be >= 1")
    _ensure_repo_path_exists()
    _resolve_repo_path(path)
    _resolve_repo_path(spdx_output_path)
    _resolve_repo_path(lint_report_path)

    missing_before = _collect_missing_spdx_headers(
        path=path,
        recursive=recursive,
        include_globs=include_globs,
        exclude_globs=exclude_globs,
        max_files=max_missing_files,
    )
    actions: list[str] = []

    if auto_fix_headers and missing_before:
        _require_mutations()
        _require_reuse_cli()
        year = str(datetime.now(timezone.utc).year)
        copyright_line = f"{year} {copyright_owner.strip() or 'Project Contributors'}"
        for chunk in _chunk_strings(missing_before, chunk_size=100):
            cmd = [
                "annotate",
                "--merge-copyrights",
                "--fallback-dot-license",
                "--license",
                default_license,
                "--copyright",
                copyright_line,
                *chunk,
            ]
            proc = _run_reuse(cmd)
            if not proc["ok"]:
                raise RuntimeError(proc["stderr"] or proc["stdout"] or "reuse annotate failed")
        actions.append(f"annotated_missing_headers:{len(missing_before)}")

    missing_after = _collect_missing_spdx_headers(
        path=path,
        recursive=recursive,
        include_globs=include_globs,
        exclude_globs=exclude_globs,
        max_files=max_missing_files,
    )

    observed_ids = _collect_spdx_license_ids(path=path, recursive=recursive)
    missing_license_texts = [
        lid
        for lid in observed_ids
        if not (REPO_PATH / "LICENSES" / f"{lid}.txt").is_file()
    ]

    if download_missing_licenses and missing_license_texts:
        _require_mutations()
        _require_reuse_cli()
        for lid in missing_license_texts:
            proc = _run_reuse(["download", lid])
            if not proc["ok"]:
                raise RuntimeError(
                    proc["stderr"] or proc["stdout"] or f"reuse download failed for {lid}"
                )
        actions.append(f"downloaded_license_texts:{len(missing_license_texts)}")
        missing_license_texts = [
            lid
            for lid in observed_ids
            if not (REPO_PATH / "LICENSES" / f"{lid}.txt").is_file()
        ]

    lint_result: dict[str, Any] | None = None
    if run_reuse_lint:
        lint_result = _run_reuse(["lint"])
        lint_abs = _resolve_repo_path(lint_report_path)
        lint_abs.parent.mkdir(parents=True, exist_ok=True)
        lint_abs.write_text(
            (
                f"reuse lint exit_code={lint_result['exit_code']}\n\n"
                f"{lint_result.get('stdout', '')}\n\n{lint_result.get('stderr', '')}\n"
            ),
            encoding="utf-8",
        )

    spdx_result: dict[str, Any] | None = None
    if generate_spdx:
        _require_reuse_cli()
        spdx_result = _run_reuse(["spdx", "-o", spdx_output_path])
        if not spdx_result["ok"]:
            raise RuntimeError(spdx_result["stderr"] or spdx_result["stdout"] or "reuse spdx failed")

    ok = (
        len(missing_after) == 0
        and len(missing_license_texts) == 0
        and (lint_result is None or bool(lint_result.get("ok", False)))
    )
    return {
        "schema": "license_monitor.v1",
        "ok": ok,
        "path": path,
        "recursive": recursive,
        "actions": actions,
        "missing_spdx_header_count": len(missing_after),
        "missing_spdx_headers": missing_after[:200],
        "observed_spdx_license_ids": observed_ids,
        "missing_license_text_count": len(missing_license_texts),
        "missing_license_texts": missing_license_texts[:200],
        "lint": lint_result,
        "spdx": spdx_result,
        "lint_report_path": lint_report_path if run_reuse_lint else None,
        "spdx_output_path": spdx_output_path if generate_spdx else None,
    }


@mcp.tool()
def install_git_hooks(
    install_pre_commit: bool = True,
    install_pre_push: bool = True,
    include_foss_reports: bool = True,
    include_lab_reports: bool = True,
    overwrite: bool = False,
) -> dict[str, Any]:
    """Install git hooks that generate FOSS and lab reports."""
    _require_mutations()
    _require_git_repo()
    hooks_rel = _git("rev-parse", "--git-path", "hooks").stdout.strip()
    hooks_path_raw = Path(hooks_rel)
    if hooks_path_raw.is_absolute():
        hooks_dir = hooks_path_raw.resolve()
        try:
            hooks_dir.relative_to(REPO_PATH)
        except ValueError as exc:
            raise ValueError("git hooks path escapes repository root") from exc
    else:
        hooks_dir = _resolve_repo_path(hooks_rel)
    hooks_dir.mkdir(parents=True, exist_ok=True)

    if not install_pre_commit and not install_pre_push:
        raise ValueError("at least one hook must be selected")

    script_lines = [
        "#!/usr/bin/env bash",
        "set -euo pipefail",
        "",
        'repo_root="$(git rev-parse --show-toplevel)"',
        'cd "$repo_root"',
        "mkdir -p .build/reports",
        "",
    ]
    if include_foss_reports:
        script_lines.extend(
            [
                "if ! command -v reuse >/dev/null 2>&1; then",
                '  echo "reuse CLI not found. Install \\"reuse\\" before committing/pushing." >&2',
                "  exit 1",
                "fi",
                "reuse lint > .build/reports/REUSE_LINT.txt",
                "reuse spdx -o .build/reports/REUSE.spdx",
                "",
            ]
        )
    script_lines.extend(
        [
            "if command -v python >/dev/null 2>&1; then",
            '  PYTHON_BIN="python"',
            "elif command -v python3 >/dev/null 2>&1; then",
            '  PYTHON_BIN="python3"',
            "else",
            '  echo "python or python3 is required for lab report hooks." >&2',
            "  exit 1",
            "fi",
            "",
        ]
    )

    pre_commit_lines = list(script_lines)
    if include_lab_reports:
        pre_commit_lines.append(
            '"$PYTHON_BIN" toolchain/dev/labs/policy_gatekeeper.py --changed-ref HEAD '
            '--report-path .build/reports/POLICY_GATEKEEPER.md'
        )

    pre_push_lines = list(script_lines)
    if include_lab_reports:
        pre_push_lines.append(
            '"$PYTHON_BIN" toolchain/dev/labs/policy_gatekeeper.py --changed-ref HEAD '
            '--report-path .build/reports/POLICY_GATEKEEPER.md'
        )
        pre_push_lines.append(
            '"$PYTHON_BIN" toolchain/dev/labs/repo_digital_twin.py '
            '--json .build/reports/REPO_DIGITAL_TWIN.json '
            '--md .build/reports/REPO_DIGITAL_TWIN.md'
        )

    installed: list[str] = []
    skipped: list[str] = []
    if install_pre_commit:
        pre_commit_path = hooks_dir / "pre-commit"
        if pre_commit_path.exists() and not overwrite:
            skipped.append(str(pre_commit_path))
        else:
            pre_commit_path.write_text("\n".join(pre_commit_lines).strip() + "\n", encoding="utf-8")
            pre_commit_path.chmod(0o755)
            installed.append(str(pre_commit_path.relative_to(REPO_PATH)))

    if install_pre_push:
        pre_push_path = hooks_dir / "pre-push"
        if pre_push_path.exists() and not overwrite:
            skipped.append(str(pre_push_path))
        else:
            pre_push_path.write_text("\n".join(pre_push_lines).strip() + "\n", encoding="utf-8")
            pre_push_path.chmod(0o755)
            installed.append(str(pre_push_path.relative_to(REPO_PATH)))

    return {
        "schema": "install_git_hooks.v1",
        "hooks_dir": str(hooks_dir.relative_to(REPO_PATH)),
        "installed": installed,
        "skipped": skipped,
        "overwrite": overwrite,
        "include_foss_reports": include_foss_reports,
        "include_lab_reports": include_lab_reports,
    }


@mcp.tool()
def find_paths(
    path: str = ".",
    recursive: bool = True,
    include_hidden: bool = False,
    include_globs: list[str] | None = None,
    exclude_globs: list[str] | None = None,
    file_type: str = "any",
    max_depth: int | None = None,
    max_entries: int = 1000,
    output_profile: str = "compact",
    offset: int = 0,
    limit: int | None = None,
    adaptive_limits: bool = True,
) -> list[str]:
    """Find files and/or directories under a repository-relative path."""
    if max_entries < 1:
        raise ValueError("max_entries must be >= 1")
    if adaptive_limits:
        max_entries = _adaptive_limit(max_entries, soft_cap=2000)
    if max_depth is not None and max_depth < 0:
        raise ValueError("max_depth must be >= 0")
    if file_type not in {"any", "file", "dir"}:
        raise ValueError("file_type must be one of: any, file, dir")
    profile = _default_output_profile(output_profile)

    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)

    results: list[str] = []

    def maybe_add(candidate: Path) -> None:
        if len(results) >= max_entries:
            return
        rel_path = candidate.relative_to(REPO_PATH)
        rel_str = str(rel_path).replace("\\", "/")

        if not include_hidden and _is_hidden_rel_path(rel_path):
            return
        if not _allowed_by_globs(rel_str, include_globs, exclude_globs):
            return

        if file_type == "file" and not candidate.is_file():
            return
        if file_type == "dir" and not candidate.is_dir():
            return

        if candidate.is_dir():
            rel_str += "/"
        results.append(rel_str)

    if root.is_file():
        maybe_add(root)
        return results

    root_parts = len(root.relative_to(REPO_PATH).parts)

    if recursive:
        iterator = root.rglob("*")
    else:
        iterator = root.glob("*")

    for item in iterator:
        if len(results) >= max_entries:
            break
        if max_depth is not None:
            item_parts = len(item.relative_to(REPO_PATH).parts)
            if item_parts - root_parts > max_depth:
                continue
        maybe_add(item)

    results.sort()
    results = _paginate(results, offset=offset, limit=limit)
    if profile == "compact":
        return [item[:-1] if item.endswith("/") else item for item in results]
    return results


@mcp.tool()
def grep(
    pattern: str,
    path: str = ".",
    recursive: bool = True,
    case_insensitive: bool = False,
    include_globs: list[str] | None = None,
    exclude_globs: list[str] | None = None,
    include_hidden: bool = False,
    max_matches: int = 500,
    max_file_bytes: int = 1048576,
    encoding: str = "utf-8",
    output_profile: str = "compact",
    fields: list[str] | None = None,
    offset: int = 0,
    limit: int | None = None,
    dedupe: bool = True,
    compress: bool = False,
    adaptive_limits: bool = True,
    summary_mode: str = "full",
    store_result: bool = False,
) -> list[dict[str, Any]]:
    """Search repository files for a regex pattern and return matches.

    Returns a list of objects: { path, line, column, match, lineText }.
    Paths are repository-relative; line/column are 1-based.
    """
    if max_matches < 1:
        raise ValueError("max_matches must be >= 1")
    if max_file_bytes < 1:
        raise ValueError("max_file_bytes must be >= 1")
    profile = _default_output_profile(output_profile)
    if summary_mode not in {"full", "quick"}:
        raise ValueError("summary_mode must be one of: full, quick")
    if adaptive_limits:
        max_matches = _adaptive_limit(max_matches, soft_cap=250)
    elif profile == "compact":
        max_matches = min(max_matches, 250)

    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)

    flags = 0
    if case_insensitive:
        flags |= re.IGNORECASE
    try:
        regex = re.compile(pattern, flags)
    except re.error as exc:
        raise ValueError(f"invalid regex pattern: {exc}") from exc

    results: list[dict[str, Any]] = []

    def search_file(p: Path) -> None:
        nonlocal results
        if not p.is_file():
            return
        rel_path = p.relative_to(REPO_PATH)
        rel_str = str(rel_path).replace("\\", "/")
        if not include_hidden and _is_hidden_rel_path(rel_path):
            return
        if not _allowed_by_globs(rel_str, include_globs, exclude_globs):
            return
        if _is_likely_binary(p, max_file_bytes=max_file_bytes):
            return

        try:
            with p.open("r", encoding=encoding, errors="replace") as f:
                for idx, line in enumerate(f, start=1):
                    for m in regex.finditer(line):
                        res = {
                            "path": rel_str,
                            "line": idx,
                            "column": m.start() + 1,
                            "match": m.group(0),
                            "lineText": line.rstrip("\n"),
                        }
                        results.append(_match_to_profile(res, profile))
                        if len(results) >= max_matches:
                            return
                    if len(results) >= max_matches:
                        return
        except OSError:
            return

    if root.is_file():
        search_file(root)
    else:
        it = root.rglob("*") if recursive else root.glob("*")
        for p in it:
            if len(results) >= max_matches:
                break
            if p.is_dir():
                continue
            search_file(p)

    if dedupe:
        uniq: dict[str, dict[str, Any]] = {}
        for row in results:
            key = f"{row.get('path')}:{row.get('line')}:{row.get('match')}"
            if key not in uniq:
                uniq[key] = row
        results = list(uniq.values())
    total = len(results)
    results = _paginate(results, offset=offset, limit=limit)
    results = _select_fields(results, fields)
    if summary_mode == "quick":
        summary = {
            "schema": "grep.quick.v1",
            "total_matches": total,
            "returned": len(results),
            "paths": sorted({str(r.get("path")) for r in results if r.get("path")})[:100],
        }
        if store_result:
            rid = _result_store_put("grep", summary)
            summary["result_id"] = rid
        return [summary]
    if compress:
        compressed = _compress_table(results)
        if store_result:
            rid = _result_store_put("grep", compressed)
            compressed["result_id"] = rid
        return [compressed]
    if store_result:
        rid = _result_store_put("grep", results)
        return [{"schema": "grep.result_handle.v1", "result_id": rid, "count": len(results)}]
    return results


@mcp.tool()
def replace_in_files(
    pattern: str,
    replacement: str,
    path: str = ".",
    recursive: bool = True,
    include_globs: list[str] | None = None,
    exclude_globs: list[str] | None = None,
    include_hidden: bool = False,
    regex: bool = True,
    case_insensitive: bool = False,
    dry_run: bool = True,
    max_files: int = 200,
    max_replacements: int = 5000,
    max_file_bytes: int = 1048576,
    encoding: str = "utf-8",
) -> dict[str, Any]:
    """Replace text in files under a path, optionally as a dry-run preview."""
    if max_files < 1:
        raise ValueError("max_files must be >= 1")
    if max_replacements < 1:
        raise ValueError("max_replacements must be >= 1")
    if max_file_bytes < 1:
        raise ValueError("max_file_bytes must be >= 1")

    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)

    if not dry_run:
        _require_mutations()

    flags = re.MULTILINE
    if case_insensitive:
        flags |= re.IGNORECASE

    if regex:
        try:
            compiled = re.compile(pattern, flags)
        except re.error as exc:
            raise ValueError(f"invalid regex pattern: {exc}") from exc
    else:
        compiled = re.compile(re.escape(pattern), flags)

    files_changed: list[dict[str, Any]] = []
    scanned_files = 0
    total_replacements = 0
    files_limit_reached = False
    replacements_limit_reached = False

    def iter_candidates():
        if root.is_file():
            yield root
            return
        if recursive:
            for p in root.rglob("*"):
                if p.is_file():
                    yield p
            return
        for p in root.glob("*"):
            if p.is_file():
                yield p

    for candidate in iter_candidates():
        if len(files_changed) >= max_files:
            files_limit_reached = True
            break

        scanned_files += 1
        rel_path = candidate.relative_to(REPO_PATH)
        rel_str = str(rel_path).replace("\\", "/")

        if not include_hidden and _is_hidden_rel_path(rel_path):
            continue
        if not _allowed_by_globs(rel_str, include_globs, exclude_globs):
            continue
        if _is_likely_binary(candidate, max_file_bytes=max_file_bytes):
            continue

        try:
            original = candidate.read_text(encoding=encoding, errors="replace")
        except OSError:
            continue

        remaining = max_replacements - total_replacements
        if remaining <= 0:
            replacements_limit_reached = True
            break

        updated, replacements = compiled.subn(replacement, original, count=remaining)
        if replacements <= 0:
            continue

        total_replacements += replacements
        files_changed.append(
            {
                "path": rel_str,
                "replacements": replacements,
                "changed": True,
            }
        )

        if not dry_run:
            candidate.write_text(updated, encoding=encoding)

        if total_replacements >= max_replacements:
            replacements_limit_reached = True
            break

    return {
        "path": str(root.relative_to(REPO_PATH)),
        "dry_run": dry_run,
        "regex": regex,
        "case_insensitive": case_insensitive,
        "scanned_files": scanned_files,
        "files_changed_count": len(files_changed),
        "total_replacements": total_replacements,
        "files_limit_reached": files_limit_reached,
        "replacements_limit_reached": replacements_limit_reached,
        "files_changed": files_changed,
    }


@mcp.tool()
def read_snippet(
    path: str,
    start_line: int,
    end_line: int,
    context_before: int = 0,
    context_after: int = 0,
    encoding: str = "utf-8",
    output_profile: str = "compact",
) -> dict[str, Any]:
    """Read a focused line range with optional surrounding context."""
    profile = _validate_output_profile(output_profile)
    file_path = _resolve_repo_path(path)
    snippet = _build_snippet(
        file_path=file_path,
        start_line=start_line,
        end_line=end_line,
        context_before=context_before,
        context_after=context_after,
        encoding=encoding,
    )
    if profile == "compact":
        return {
            "path": snippet["path"],
            "start_line": snippet["start_line"],
            "end_line": snippet["end_line"],
            "content": snippet["content"],
        }
    return snippet


@mcp.tool()
def read_batch(
    requests: list[dict[str, Any]],
    encoding: str = "utf-8",
    max_items: int = 50,
    output_profile: str = "compact",
) -> dict[str, Any]:
    """Read multiple focused snippets in one call."""
    if max_items < 1:
        raise ValueError("max_items must be >= 1")
    profile = _validate_output_profile(output_profile)
    if len(requests) > max_items:
        raise ValueError(f"too many requests; max_items={max_items}")

    snippets: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []

    for idx, req in enumerate(requests, start=1):
        path = req.get("path")
        start_line = req.get("start_line")
        end_line = req.get("end_line")
        context_before = int(req.get("context_before", 0))
        context_after = int(req.get("context_after", 0))
        if not isinstance(path, str) or not isinstance(start_line, int) or not isinstance(
            end_line, int
        ):
            errors.append({"index": idx, "error": "path/start_line/end_line are required"})
            continue
        try:
            snippet = _build_snippet(
                file_path=_resolve_repo_path(path),
                start_line=start_line,
                end_line=end_line,
                context_before=context_before,
                context_after=context_after,
                encoding=encoding,
            )
            if profile == "compact":
                snippet = {
                    "path": snippet["path"],
                    "start_line": snippet["start_line"],
                    "end_line": snippet["end_line"],
                    "content": snippet["content"],
                }
            snippets.append(snippet)
        except Exception as exc:
            errors.append({"index": idx, "path": path, "error": str(exc)})

    return {
        "count": len(snippets),
        "error_count": len(errors),
        "snippets": snippets,
        "errors": errors if profile != "compact" else errors[:10],
    }


@mcp.tool()
def semantic_find(
    query: str,
    path: str = ".",
    max_results: int = 30,
    output_profile: str | None = None,
    include_private_symbols: bool = False,
    fields: list[str] | None = None,
    offset: int = 0,
    limit: int | None = None,
    dedupe: bool = True,
    compress: bool = False,
    adaptive_limits: bool = True,
    summary_mode: str = "full",
    store_result: bool = False,
    use_local_rerank: bool = False,
    local_rerank_top_k: int = 50,
) -> dict[str, Any]:
    """Ranked search over file paths, symbols, and text matches."""
    if not query.strip():
        raise ValueError("query must not be empty")
    if max_results < 1:
        raise ValueError("max_results must be >= 1")
    if summary_mode not in {"full", "quick"}:
        raise ValueError("summary_mode must be one of: full, quick")
    if adaptive_limits:
        max_results = _adaptive_limit(max_results, soft_cap=100)
    profile = _default_output_profile(output_profile)
    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)

    terms = [t.lower() for t in re.split(r"\s+", query.strip()) if t]
    candidates: dict[str, dict[str, Any]] = {}

    for rel in find_paths(
        path=path,
        recursive=True,
        include_hidden=False,
        max_entries=5000,
        output_profile="compact",
        file_type="file",
    ):
        score = 0.0
        rel_low = rel.lower()
        for term in terms:
            if term in rel_low:
                score += 3.0
        if score > 0:
            candidates.setdefault(
                f"path:{rel}",
                {"kind": "path", "path": rel, "score": 0.0, "reasons": []},
            )
            candidates[f"path:{rel}"]["score"] += score
            candidates[f"path:{rel}"]["reasons"].append("path_term_match")

    symbol_term = "|".join(re.escape(t) for t in terms if t)
    if symbol_term:
        for sym in symbol_index(
            path=path,
            include_private=include_private_symbols,
            recursive=True,
            max_symbols=5000,
            output_profile="normal",
        ):
            name = str(sym.get("name", ""))
            score = 0.0
            for term in terms:
                if term in name.lower():
                    score += 5.0
            if score <= 0:
                continue
            key = f"symbol:{sym['path']}:{name}:{sym['line_start']}"
            candidates[key] = {
                "kind": "symbol",
                "path": sym["path"],
                "name": name,
                "line_start": sym["line_start"],
                "line_end": sym.get("line_end"),
                "score": score,
                "reasons": ["symbol_name_match"],
            }

    pattern = "|".join(re.escape(t) for t in terms)
    if pattern:
        matches = grep(
            pattern=pattern,
            path=path,
            recursive=True,
            case_insensitive=True,
            max_matches=200,
            output_profile="compact",
        )
        for m in matches:
            key = f"grep:{m['path']}:{m['line']}:{m['column']}"
            candidates[key] = {
                "kind": "text_match",
                "path": m["path"],
                "line": m["line"],
                "column": m["column"],
                "match": m["match"],
                "score": 2.0,
                "reasons": ["text_match"],
            }

    ranked = sorted(candidates.values(), key=lambda x: x["score"], reverse=True)
    if dedupe:
        uniq: dict[str, dict[str, Any]] = {}
        for row in ranked:
            key = f"{row.get('kind')}:{row.get('path')}:{row.get('line', '')}:{row.get('name', '')}"
            if key not in uniq:
                uniq[key] = row
        ranked = list(uniq.values())
    ranked = ranked[:max_results]
    if use_local_rerank and ranked:
        reranked = local_rerank(
            query=query,
            candidates=ranked[: max(local_rerank_top_k, max_results)],
            top_k=max_results,
            backend="auto",
            output_profile="normal",
        )
        if isinstance(reranked, dict) and isinstance(reranked.get("results"), list):
            ranked = reranked["results"]
    ranked = _paginate(ranked, offset=offset, limit=limit)
    if profile == "compact":
        ranked = [
            {
                "kind": r.get("kind"),
                "path": r.get("path"),
                "score": r.get("score"),
            }
            for r in ranked
        ]
    ranked = _select_fields(ranked, fields)
    result = {
        "schema": "semantic_find.v1",
        "query": query,
        "path": str(root.relative_to(REPO_PATH)),
        "count": len(ranked),
        "results": ranked,
    }
    if summary_mode == "quick":
        result = {
            "schema": "semantic_find.quick.v1",
            "query": query,
            "count": len(ranked),
            "top_paths": [r.get("path") for r in ranked[:20] if isinstance(r, dict)],
        }
    if compress and isinstance(result.get("results"), list):
        result["results_compressed"] = _compress_table(result["results"])
        result.pop("results", None)
    if store_result:
        rid = _result_store_put("semantic_find", result)
        result["result_id"] = rid
    return result


@mcp.tool()
def symbol_index(
    path: str = ".",
    include_private: bool = False,
    recursive: bool = True,
    max_symbols: int = 5000,
    encoding: str = "utf-8",
    output_profile: str = "compact",
    fields: list[str] | None = None,
    offset: int = 0,
    limit: int | None = None,
    adaptive_limits: bool = True,
) -> list[dict[str, Any]]:
    """Index Python symbols (classes/functions) for focused navigation."""
    if max_symbols < 1:
        raise ValueError("max_symbols must be >= 1")
    if adaptive_limits:
        max_symbols = _adaptive_limit(max_symbols, soft_cap=2000)
    profile = _default_output_profile(output_profile)
    if profile == "compact":
        max_symbols = min(max_symbols, 2000)

    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)

    fingerprint = _fingerprint_path(
        root, recursive=recursive, suffixes={".py"}, max_files=3000
    )
    cache_key = json.dumps(
        {
            "path": str(root.relative_to(REPO_PATH)),
            "include_private": include_private,
            "recursive": recursive,
            "encoding": encoding,
            "fingerprint": fingerprint,
        },
        sort_keys=True,
    )
    cached = _cache_get("symbol_index", cache_key)
    if isinstance(cached, list):
        symbols = [dict(row) for row in cached]
    else:
        symbols = []
        for candidate in _iter_candidate_files(root, recursive=recursive):
            if candidate.suffix != ".py":
                continue
            rel_path = candidate.relative_to(REPO_PATH)
            rel_str = str(rel_path).replace("\\", "/")
            if _is_likely_binary(candidate):
                continue
            try:
                source = candidate.read_text(encoding=encoding, errors="replace")
            except OSError:
                continue
            extracted = _collect_python_symbols(
                source, rel_str, include_private=include_private
            )
            symbols.extend(extracted)
        _cache_set("symbol_index", cache_key, symbols)

    profiled = [_symbol_to_profile(symbol, profile) for symbol in symbols]
    profiled = profiled[:max_symbols]
    profiled = _paginate(profiled, offset=offset, limit=limit)
    profiled = _select_fields(profiled, fields)
    return profiled


@mcp.tool()
def dependency_map(
    path: str = ".",
    recursive: bool = True,
    include_stdlib: bool = False,
    max_files: int = 3000,
    output_profile: str = "compact",
    encoding: str = "utf-8",
    fields: list[str] | None = None,
    offset: int = 0,
    limit: int | None = None,
    adaptive_limits: bool = True,
    summary_mode: str = "full",
    compress: bool = False,
    store_result: bool = False,
) -> dict[str, Any]:
    """Build a Python import dependency map for repo-local modules."""
    if max_files < 1:
        raise ValueError("max_files must be >= 1")
    if summary_mode not in {"full", "quick"}:
        raise ValueError("summary_mode must be one of: full, quick")
    if adaptive_limits:
        max_files = _adaptive_limit(max_files, soft_cap=1500)
    profile = _default_output_profile(output_profile)
    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)

    fingerprint = _fingerprint_path(
        root, recursive=recursive, suffixes={".py"}, max_files=max_files
    )
    cache_key = json.dumps(
        {
            "path": str(root.relative_to(REPO_PATH)),
            "recursive": recursive,
            "include_stdlib": include_stdlib,
            "encoding": encoding,
            "fingerprint": fingerprint,
            "max_files": max_files,
        },
        sort_keys=True,
    )
    cached = _cache_get("dependency_map", cache_key)
    if isinstance(cached, dict):
        python_file_count = int(cached.get("python_file_count", 0))
        edges = list(cached.get("edges", []))
        unresolved = list(cached.get("unresolved_imports", []))
    else:
        module_to_path: dict[str, str] = {}
        python_files: list[Path] = []
        for candidate in _iter_candidate_files(root, recursive=recursive):
            if candidate.suffix != ".py":
                continue
            rel = candidate.relative_to(REPO_PATH)
            module = _module_name_from_relpath(rel)
            rel_str = str(rel).replace("\\", "/")
            module_to_path[module] = rel_str
            python_files.append(candidate)
            if len(python_files) >= max_files:
                break

        edges: list[dict[str, Any]] = []
        unresolved: list[dict[str, Any]] = []

        for file_path in python_files:
            rel = str(file_path.relative_to(REPO_PATH)).replace("\\", "/")
            try:
                tree = ast.parse(file_path.read_text(encoding=encoding, errors="replace"))
            except (OSError, SyntaxError):
                continue
            for node in ast.walk(tree):
                if isinstance(node, ast.Import):
                    imports = [alias.name for alias in node.names]
                elif isinstance(node, ast.ImportFrom):
                    if node.module:
                        imports = [node.module]
                    else:
                        imports = []
                else:
                    continue

                for imp in imports:
                    resolved_path = None
                    for candidate_module in _import_candidates(imp):
                        resolved = module_to_path.get(candidate_module)
                        if resolved:
                            resolved_path = resolved
                            break
                    if resolved_path:
                        edges.append(
                            {
                                "from": rel,
                                "to": resolved_path,
                                "import": imp,
                                "line": int(getattr(node, "lineno", 1)),
                            }
                        )
                    elif include_stdlib:
                        unresolved.append(
                            {
                                "from": rel,
                                "import": imp,
                                "line": int(getattr(node, "lineno", 1)),
                            }
                        )
        python_file_count = len(python_files)
        _cache_set(
            "dependency_map",
            cache_key,
            {
                "python_file_count": python_file_count,
                "edges": edges,
                "unresolved_imports": unresolved,
            },
        )

    result: dict[str, Any] = {
        "schema": "dependency_map.v1",
        "root": str(root.relative_to(REPO_PATH)),
        "python_file_count": python_file_count,
        "edge_count": len(edges),
        "edges": _select_fields(_paginate(edges, offset=offset, limit=limit), fields),
    }
    if profile == "compact":
        compact = {
            "schema": "dependency_map.compact.v1",
            "python_file_count": result["python_file_count"],
            "edge_count": result["edge_count"],
            "edges": result["edges"][:500] if isinstance(result["edges"], list) else [],
        }
        if compress:
            compact["edges_compressed"] = _compress_table(compact["edges"])
            compact.pop("edges", None)
        if store_result:
            compact["result_id"] = _result_store_put("dependency_map", compact)
        return compact
    if profile == "verbose":
        result["unresolved_imports"] = unresolved
        inbound: dict[str, int] = {}
        outbound: dict[str, int] = {}
        for edge in edges:
            inbound[edge["to"]] = inbound.get(edge["to"], 0) + 1
            outbound[edge["from"]] = outbound.get(edge["from"], 0) + 1
        result["hotspots"] = {
            "most_imported": sorted(
                [{"path": k, "count": v} for k, v in inbound.items()],
                key=lambda x: x["count"],
                reverse=True,
            )[:20],
            "most_importing": sorted(
                [{"path": k, "count": v} for k, v in outbound.items()],
                key=lambda x: x["count"],
                reverse=True,
            )[:20],
        }
    if summary_mode == "quick":
        result = {
            "schema": "dependency_map.quick.v1",
            "root": result["root"],
            "python_file_count": result["python_file_count"],
            "edge_count": result["edge_count"],
        }
    if compress and isinstance(result.get("edges"), list):
        result["edges_compressed"] = _compress_table(result["edges"])
        result.pop("edges", None)
    if store_result:
        result["result_id"] = _result_store_put("dependency_map", result)
    return result


@mcp.tool()
def call_graph(
    path: str = ".",
    recursive: bool = True,
    max_edges: int = 5000,
    output_profile: str | None = None,
    encoding: str = "utf-8",
    fields: list[str] | None = None,
    offset: int = 0,
    limit: int | None = None,
    adaptive_limits: bool = True,
    summary_mode: str = "full",
    compress: bool = False,
    store_result: bool = False,
) -> dict[str, Any]:
    """Build a simple Python function-level call graph."""
    if max_edges < 1:
        raise ValueError("max_edges must be >= 1")
    if summary_mode not in {"full", "quick"}:
        raise ValueError("summary_mode must be one of: full, quick")
    if adaptive_limits:
        max_edges = _adaptive_limit(max_edges, soft_cap=2000)
    profile = _default_output_profile(output_profile)
    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)

    fingerprint = _fingerprint_path(
        root, recursive=recursive, suffixes={".py"}, max_files=3000
    )
    cache_key = json.dumps(
        {
            "path": str(root.relative_to(REPO_PATH)),
            "recursive": recursive,
            "encoding": encoding,
            "fingerprint": fingerprint,
            "max_edges": max_edges,
        },
        sort_keys=True,
    )
    cached = _cache_get("call_graph", cache_key)
    if isinstance(cached, list):
        edges = list(cached)
    else:
        edges: list[dict[str, Any]] = []
        for candidate in _iter_candidate_files(root, recursive=recursive):
            if candidate.suffix != ".py":
                continue
            rel = str(candidate.relative_to(REPO_PATH)).replace("\\", "/")
            try:
                tree = ast.parse(candidate.read_text(encoding=encoding, errors="replace"))
            except (SyntaxError, OSError):
                continue
            for node in ast.walk(tree):
                if not isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                    continue
                caller = node.name
                for child in ast.walk(node):
                    if isinstance(child, ast.Call):
                        callee = _ast_expr_name(child.func)
                        if not callee:
                            continue
                        edges.append(
                            {
                                "path": rel,
                                "caller": caller,
                                "callee": callee,
                                "line": int(getattr(child, "lineno", getattr(node, "lineno", 1))),
                            }
                        )
                        if len(edges) >= max_edges:
                            break
                if len(edges) >= max_edges:
                    break
            if len(edges) >= max_edges:
                break
        _cache_set("call_graph", cache_key, edges)

    paged_edges = _select_fields(_paginate(edges, offset=offset, limit=limit), fields)

    if profile == "compact":
        compact = {"schema": "call_graph.compact.v1", "edge_count": len(edges), "edges": paged_edges[:500]}
        if compress:
            compact["edges_compressed"] = _compress_table(compact["edges"])
            compact.pop("edges", None)
        if store_result:
            compact["result_id"] = _result_store_put("call_graph", compact)
        return compact
    inbound: dict[str, int] = {}
    for edge in edges:
        inbound[edge["callee"]] = inbound.get(edge["callee"], 0) + 1
    result: dict[str, Any] = {"schema": "call_graph.v1", "edge_count": len(edges), "edges": paged_edges}
    if profile == "verbose":
        result["most_called"] = sorted(
            [{"symbol": k, "count": v} for k, v in inbound.items()],
            key=lambda x: x["count"],
            reverse=True,
        )[:25]
    if summary_mode == "quick":
        result = {
            "schema": "call_graph.quick.v1",
            "edge_count": result["edge_count"],
            "top_called": result.get("most_called", [])[:10] if profile == "verbose" else [],
        }
    if compress and isinstance(result.get("edges"), list):
        result["edges_compressed"] = _compress_table(result["edges"])
        result.pop("edges", None)
    if store_result:
        result["result_id"] = _result_store_put("call_graph", result)
    return result


@mcp.tool()
def ast_search(
    path: str = ".",
    node_type: str = "Call",
    name_pattern: str | None = None,
    recursive: bool = True,
    max_results: int = 500,
    encoding: str = "utf-8",
) -> list[dict[str, Any]]:
    """Search Python AST nodes to find structural code matches."""
    if max_results < 1:
        raise ValueError("max_results must be >= 1")
    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)

    try:
        node_cls = getattr(ast, node_type)
    except AttributeError as exc:
        raise ValueError(f"unsupported node_type: {node_type}") from exc

    name_regex = re.compile(name_pattern) if name_pattern else None
    results: list[dict[str, Any]] = []

    for candidate in _iter_candidate_files(root, recursive=recursive):
        if candidate.suffix != ".py":
            continue
        rel_str = str(candidate.relative_to(REPO_PATH)).replace("\\", "/")
        try:
            source = candidate.read_text(encoding=encoding, errors="replace")
            tree = ast.parse(source)
        except (OSError, SyntaxError):
            continue

        for node in ast.walk(tree):
            if not isinstance(node, node_cls):
                continue
            node_name = _node_display_name(node)
            if name_regex and not name_regex.search(node_name):
                continue
            results.append(
                {
                    "path": rel_str,
                    "node_type": node_type,
                    "name": node_name,
                    "line": int(getattr(node, "lineno", 1)),
                    "column": int(getattr(node, "col_offset", 0)) + 1,
                    "end_line": int(getattr(node, "end_lineno", getattr(node, "lineno", 1))),
                }
            )
            if len(results) >= max_results:
                return results
    return results


@mcp.tool()
def apply_unified_diff(
    diff_text: str,
    check_only: bool = True,
    cached: bool = False,
) -> dict[str, Any]:
    """Apply a unified diff through git-apply with optional dry-run checks."""
    _require_git_repo()
    if not check_only:
        _require_mutations()

    args = ["git", "-C", str(REPO_PATH), "apply"]
    if check_only:
        args.append("--check")
    if cached:
        args.append("--cached")

    proc = subprocess.run(
        args,
        input=diff_text,
        check=False,
        capture_output=True,
        text=True,
    )
    return {
        "ok": proc.returncode == 0,
        "check_only": check_only,
        "cached": cached,
        "exit_code": proc.returncode,
        "stdout": _trim_text(proc.stdout.strip()),
        "stderr": _trim_text(proc.stderr.strip()),
    }


@mcp.tool()
def command_runner(
    command: list[str],
    cwd: str = ".",
    timeout_seconds: int = 30,
    max_output_chars: int | None = None,
) -> dict[str, Any]:
    """Run a whitelisted command without shell interpolation."""
    if not command:
        raise ValueError("command must not be empty")
    if timeout_seconds < 1:
        raise ValueError("timeout_seconds must be >= 1")
    out_cap = _token_budget_apply_max(max_output_chars)

    binary = command[0]
    if binary not in SAFE_COMMANDS:
        raise ValueError(f"command not allowed: {binary}")

    if binary == "git":
        if len(command) < 2:
            raise ValueError("git command must include a subcommand")
        if command[1] not in SAFE_GIT_SUBCOMMANDS:
            raise ValueError(f"git subcommand not allowed: {command[1]}")
    if binary == "sed" and any(arg == "-i" or arg.startswith("-i") for arg in command[1:]):
        raise ValueError("sed in-place edits are not allowed")
    if binary == "find" and any(arg in {"-delete", "-exec", "-ok"} for arg in command[1:]):
        raise ValueError("find destructive/exec flags are not allowed")
    if binary == "awk":
        script = command[1] if len(command) > 1 else ""
        if "system(" in script:
            raise ValueError("awk system() is not allowed")

    workdir = _resolve_repo_path(cwd)
    try:
        proc = subprocess.run(
            command,
            cwd=str(workdir),
            check=False,
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
        )
    except subprocess.TimeoutExpired as exc:
        _failure_record(
            command=command,
            stderr="command timed out",
            stdout=(exc.stdout or "") if isinstance(exc.stdout, str) else "",
            category="command_runner",
            suggestion="Increase timeout_seconds or narrow command scope.",
        )
        return {
            "ok": False,
            "exit_code": None,
            "command": command,
            "cwd": str(workdir.relative_to(REPO_PATH)),
            "stdout": _trim_text((exc.stdout or "") if isinstance(exc.stdout, str) else "", max_chars=out_cap),
            "stderr": _trim_text((exc.stderr or "") if isinstance(exc.stderr, str) else "", max_chars=out_cap),
            "timeout": True,
        }
    except FileNotFoundError as exc:
        _failure_record(
            command=command,
            stderr=str(exc),
            category="command_runner",
            suggestion="Verify the executable is installed in the runtime.",
        )
        return {
            "ok": False,
            "exit_code": None,
            "command": command,
            "cwd": str(workdir.relative_to(REPO_PATH)),
            "stdout": "",
            "stderr": str(exc),
            "timeout": False,
        }
    if proc.returncode != 0:
        _failure_record(
            command=command,
            stderr=proc.stderr,
            stdout=proc.stdout,
            category="command_runner",
            suggestion="Inspect stderr and retry with narrower scope or valid flags.",
        )
    return {
        "ok": proc.returncode == 0,
        "exit_code": proc.returncode,
        "command": command,
        "cwd": str(workdir.relative_to(REPO_PATH)),
        "stdout": _trim_text(proc.stdout, max_chars=out_cap),
        "stderr": _trim_text(proc.stderr, max_chars=out_cap),
    }


@mcp.tool()
def summarize_diff(
    ref: str | None = None,
    staged: bool = False,
    pathspec: str | None = None,
    output_profile: str = "compact",
    include_patch: bool = False,
    patch_unified: int = 0,
) -> dict[str, Any]:
    """Return compact structured diff summary with risk hints."""
    _require_git_repo()
    profile = _validate_output_profile(output_profile)
    if patch_unified < 0:
        raise ValueError("patch_unified must be >= 0")
    args = ["diff"]
    if staged:
        args.append("--staged")
    if ref:
        args.append(ref)
    if pathspec:
        _resolve_repo_path(pathspec)
        args.extend(["--", pathspec])

    numstat = _git(*args, "--numstat").stdout
    patch = _git(*args, f"--unified={patch_unified}").stdout

    files: list[dict[str, Any]] = []
    total_added = 0
    total_deleted = 0
    risky_files: list[str] = []

    for line in numstat.splitlines():
        parts = line.split("\t")
        if len(parts) != 3:
            continue
        add_raw, del_raw, file_path = parts
        added = int(add_raw) if add_raw.isdigit() else 0
        deleted = int(del_raw) if del_raw.isdigit() else 0
        total_added += added
        total_deleted += deleted
        files.append({"path": file_path, "added": added, "deleted": deleted})
        low = file_path.lower()
        if (
            "dockerfile" in low
            or "requirements" in low
            or low.endswith(".lock")
            or low.endswith("package.json")
            or "/.github/workflows/" in low
        ):
            risky_files.append(file_path)

    todo_hits = 0
    for line in patch.splitlines():
        if not line.startswith("+") or line.startswith("+++"):
            continue
        if "TODO" in line or "FIXME" in line or "XXX" in line:
            todo_hits += 1

    result = {
        "file_count": len(files),
        "total_added": total_added,
        "total_deleted": total_deleted,
        "files": files,
        "risk_flags": {
            "risky_files": risky_files,
            "todo_like_additions": todo_hits,
        },
    }
    if profile == "compact":
        return {
            "file_count": result["file_count"],
            "total_added": result["total_added"],
            "total_deleted": result["total_deleted"],
            "risk_flags": result["risk_flags"],
        }
    if profile == "verbose":
        result["files_sorted_by_churn"] = sorted(
            files, key=lambda x: x["added"] + x["deleted"], reverse=True
        )
    if include_patch:
        result["patch"] = _trim_text(patch)
        result["patch_unified"] = patch_unified
    return result


@mcp.tool()
def json_query(
    path: str,
    query: str = "",
    file_type: str | None = None,
    output_profile: str = "compact",
) -> dict[str, Any]:
    """Query JSON/TOML/YAML content with a dot/index path."""
    profile = _validate_output_profile(output_profile)
    file_path = _resolve_repo_path(path)
    if not file_path.is_file():
        raise FileNotFoundError(path)

    fmt = (file_type or _guess_file_type(file_path)).lower()
    raw = file_path.read_text(encoding="utf-8", errors="replace")

    if fmt == "json":
        data = json.loads(raw)
    elif fmt == "toml":
        if tomllib is None:
            raise RuntimeError("tomllib is not available in this Python runtime")
        data = tomllib.loads(raw)
    elif fmt == "yaml":
        if yaml is None:
            raise RuntimeError("PyYAML is not installed in this runtime")
        data = yaml.safe_load(raw)
    else:
        raise ValueError("file_type must be one of: json, toml, yaml")

    value = _query_value(data, query) if query.strip() else data
    encoded = json.dumps(value, indent=2, ensure_ascii=True)
    result = {
        "path": str(file_path.relative_to(REPO_PATH)),
        "query": query,
        "file_type": fmt,
        "value": value,
        "value_json": _trim_text(encoded),
    }
    if profile == "compact":
        return {
            "path": result["path"],
            "query": result["query"],
            "file_type": result["file_type"],
            "value_json": result["value_json"],
        }
    if profile == "verbose":
        result["value_type"] = type(value).__name__
    return result


@mcp.tool()
def prompt_optimize(
    prompt: str,
    mode: str = "coding",
    max_chars: int = 2000,
) -> dict[str, Any]:
    """Produce a compact prompt variant tuned for low-token tool workflows."""
    if not prompt.strip():
        raise ValueError("prompt must not be empty")
    if max_chars < 100:
        raise ValueError("max_chars must be >= 100")
    if mode not in {"coding", "review", "search"}:
        raise ValueError("mode must be one of: coding, review, search")

    header = {
        "coding": "Goal: implement minimal safe change. Use compact outputs and bounded queries.",
        "review": "Goal: find high-severity issues first. Return concise findings with file/line.",
        "search": "Goal: locate exact targets quickly. Use fields, pagination, and result handles.",
    }[mode]
    suffix = (
        "Constraints: prefer output_profile=compact; set fields; use offset/limit; "
        "use summary_mode=quick first; store_result=true for large outputs."
    )
    body = re.sub(r"\s+", " ", prompt.strip())
    optimized = f"{header} Request: {body} {suffix}".strip()
    if len(optimized) > max_chars:
        optimized = optimized[:max_chars]
    return {
        "schema": "prompt_optimize.v1",
        "mode": mode,
        "original_chars": len(prompt),
        "optimized_chars": len(optimized),
        "optimized_prompt": optimized,
    }


@mcp.tool()
def math_parser(
    text: str,
    symbols: list[str] | None = None,
) -> dict[str, Any]:
    """Parse math expression text into a canonical symbolic form."""
    _require_sympy()
    if not text.strip():
        raise ValueError("text must not be empty")
    local_symbols = {}
    for s in symbols or []:
        local_symbols[s] = sp.symbols(s)
    expr = sp.sympify(text, locals=local_symbols)
    return {
        "schema": "math_parser.v1",
        "input": text,
        "parsed": str(expr),
        "latex": sp.latex(expr),
        "free_symbols": sorted(str(s) for s in expr.free_symbols),
    }


@mcp.tool()
def math_solver(
    mode: str = "simplify",
    expression: str = "",
    variable: str = "x",
    equations: list[str] | None = None,
    matrix_a: list[list[float]] | None = None,
    matrix_b: list[list[float]] | None = None,
    assumptions: dict[str, str] | None = None,
    include_steps: bool = True,
) -> dict[str, Any]:
    """Offline symbolic math solver with exact + numeric outputs."""
    _require_sympy()
    if mode not in {"simplify", "solve", "differentiate", "integrate", "matrix", "optimize"}:
        raise ValueError("mode must be one of: simplify, solve, differentiate, integrate, matrix, optimize")
    x = sp.symbols(variable, **(assumptions or {}))
    result: dict[str, Any] = {"schema": "math_solver.v1", "mode": mode}

    if mode == "simplify":
        expr = _math_expr(expression)
        exact = sp.simplify(expr)
        result["exact"] = str(exact)
        result["numeric"] = str(sp.N(exact))
    elif mode == "solve":
        eqs = equations or ([expression] if expression else [])
        if not eqs:
            raise ValueError("expression or equations required for solve mode")
        parsed = []
        for e in eqs:
            if "=" in e:
                left, right = e.split("=", 1)
                parsed.append(sp.Eq(sp.sympify(left), sp.sympify(right)))
            else:
                parsed.append(sp.Eq(sp.sympify(e), 0))
        sols = sp.solve(parsed, x, dict=True)
        result["solutions"] = [{str(k): str(v) for k, v in row.items()} for row in sols]
    elif mode == "differentiate":
        expr = _math_expr(expression)
        deriv = sp.diff(expr, x)
        result["exact"] = str(deriv)
        result["numeric"] = str(sp.N(deriv))
    elif mode == "integrate":
        expr = _math_expr(expression)
        integ = sp.integrate(expr, x)
        result["exact"] = str(integ)
        result["numeric"] = str(sp.N(integ))
    elif mode == "matrix":
        if matrix_a is None:
            raise ValueError("matrix_a is required for matrix mode")
        A = sp.Matrix(matrix_a)
        result["shape"] = [int(A.rows), int(A.cols)]
        result["determinant"] = str(A.det()) if A.rows == A.cols else None
        result["rank"] = int(A.rank())
        if matrix_b is not None:
            B = sp.Matrix(matrix_b)
            result["product"] = [[str(v) for v in row] for row in (A * B).tolist()]
    else:  # optimize
        expr = _math_expr(expression)
        deriv = sp.diff(expr, x)
        critical = sp.solve(sp.Eq(deriv, 0), x)
        result["critical_points"] = [str(v) for v in critical]
        result["derivative"] = str(deriv)

    if include_steps:
        result["steps"] = _math_steps_stub(mode, expression or str(equations or ""))
    return result


@mcp.tool()
def math_verify(
    left: str,
    right: str,
    variables: list[str] | None = None,
    trials: int = 5,
) -> dict[str, Any]:
    """Verify algebraic identity/equality via symbolic simplification and sampling."""
    _require_sympy()
    if trials < 1:
        raise ValueError("trials must be >= 1")
    l = _math_expr(left)
    r = _math_expr(right)
    diff = sp.simplify(l - r)
    proven = diff == 0
    checks: list[dict[str, Any]] = []
    syms = sorted(diff.free_symbols, key=lambda s: str(s))
    if variables:
        syms = [sp.symbols(v) for v in variables]
    if not proven and syms:
        for i in range(trials):
            vals = {s: (i + 2) for s in syms}
            ok = sp.simplify((l - r).subs(vals)) == 0
            checks.append({"substitution": {str(k): str(v) for k, v in vals.items()}, "ok": bool(ok)})
    return {
        "schema": "math_verify.v1",
        "left": left,
        "right": right,
        "proven": bool(proven),
        "difference": str(diff),
        "checks": checks,
    }


@mcp.tool()
def sql_expert(
    mode: str = "format",
    query: str = "",
    dialect: str = "generic",
    nl_request: str = "",
) -> dict[str, Any]:
    """Offline SQL helper for formatting, linting, and NL-to-SQL skeletons."""
    if mode not in {"format", "lint", "nl2sql"}:
        raise ValueError("mode must be one of: format, lint, nl2sql")
    result: dict[str, Any] = {"schema": "sql_expert.v1", "mode": mode, "dialect": dialect}
    if mode == "format":
        if not query.strip():
            raise ValueError("query must not be empty")
        result["formatted"] = _sql_normalize(query)
        return result
    if mode == "lint":
        if not query.strip():
            raise ValueError("query must not be empty")
        issues: list[str] = []
        q = query.lower()
        if "select *" in q:
            issues.append("Avoid SELECT * in production queries.")
        if " where " not in q and (" update " in q or " delete " in q):
            issues.append("UPDATE/DELETE without WHERE can affect all rows.")
        if " order by " in q and " limit " not in q:
            issues.append("Consider LIMIT when ORDER BY is used in high-cardinality tables.")
        result["issues"] = issues
        result["formatted"] = _sql_normalize(query)
        return result
    # nl2sql
    req = nl_request.strip().lower()
    if not req:
        raise ValueError("nl_request must not be empty for nl2sql mode")
    table = "items"
    if "user" in req:
        table = "users"
    elif "order" in req:
        table = "orders"
    skeleton = f"SELECT id, * FROM {table} WHERE <condition> ORDER BY id DESC LIMIT 100;"
    result["sql_skeleton"] = skeleton
    return result


@mcp.tool()
def security_triage(
    diff_text: str = "",
    paths: list[str] | None = None,
    max_findings: int = 100,
) -> dict[str, Any]:
    """Classify security-sensitive changes from diff snippets and path heuristics."""
    if max_findings < 1:
        raise ValueError("max_findings must be >= 1")
    findings: list[dict[str, Any]] = []
    lines = _extract_diff_lines(diff_text)
    lower_paths = [p.lower() for p in (paths or [])]
    patterns = [
        ("hardcoded_secret", re.compile(r"(api[_-]?key|secret|token)\s*[:=]\s*['\"][^'\"]+['\"]", re.IGNORECASE), "high"),
        ("command_injection", re.compile(r"(subprocess\.|os\.system|eval\()", re.IGNORECASE), "high"),
        ("sql_injection", re.compile(r"(select .* \+|f\"select|execute\(.+\+)", re.IGNORECASE), "high"),
        ("weak_crypto", re.compile(r"\b(md5|sha1)\b", re.IGNORECASE), "medium"),
    ]
    for ln in lines:
        for rule, rx, sev in patterns:
            if rx.search(ln):
                findings.append({"rule": rule, "severity": sev, "line": ln[:300]})
                if len(findings) >= max_findings:
                    break
        if len(findings) >= max_findings:
            break
    if any("auth" in p or "security" in p or "crypto" in p for p in lower_paths):
        findings.append({"rule": "sensitive_path", "severity": "medium", "line": "Sensitive path changed"})
    sev_rank = {"critical": 4, "high": 3, "medium": 2, "low": 1}
    top = "low"
    if findings:
        top = max((f["severity"] for f in findings), key=lambda s: sev_rank.get(s, 0))
    return {
        "schema": "security_triage.v1",
        "finding_count": len(findings),
        "top_severity": top,
        "findings": findings[:max_findings],
    }


@mcp.tool()
def doc_summarizer_small(
    text: str,
    max_bullets: int = 8,
    max_chars: int = 1200,
) -> dict[str, Any]:
    """Small offline summarizer for logs/docs using sentence ranking."""
    if max_bullets < 1:
        raise ValueError("max_bullets must be >= 1")
    sentences = re.split(r"(?<=[.!?])\s+", " ".join(text.split()))
    scores = []
    for s in sentences:
        if not s.strip():
            continue
        score = len(re.findall(r"\b(error|fail|warning|todo|fix|critical|security)\b", s, re.IGNORECASE))
        score += min(4, len(s) // 80)
        scores.append((score, s.strip()))
    scores.sort(key=lambda x: x[0], reverse=True)
    bullets = [s for _, s in scores[:max_bullets]]
    summary = "\n".join(f"- {b}" for b in bullets)
    return {
        "schema": "doc_summarizer_small.v1",
        "bullet_count": len(bullets),
        "summary": _trim_text(summary, max_chars=max_chars),
    }


@mcp.tool()
def code_review_classifier(
    findings: list[dict[str, Any]],
    include_confidence: bool = True,
) -> dict[str, Any]:
    """Classify review findings into bug/perf/style/security buckets."""
    buckets = {"bug": [], "perf": [], "style": [], "security": [], "other": []}
    for f in findings:
        text = " ".join(str(f.get(k, "")) for k in ("title", "message", "detail", "rule")).lower()
        bucket = "other"
        if any(k in text for k in ("injection", "secret", "xss", "auth", "csrf", "crypto")):
            bucket = "security"
        elif any(k in text for k in ("crash", "exception", "null", "wrong", "bug", "incorrect")):
            bucket = "bug"
        elif any(k in text for k in ("slow", "n+1", "latency", "optimize", "allocation")):
            bucket = "perf"
        elif any(k in text for k in ("format", "naming", "style", "lint", "readability")):
            bucket = "style"
        row = dict(f)
        if include_confidence:
            row["confidence"] = 0.8 if bucket != "other" else 0.5
        buckets[bucket].append(row)
    return {
        "schema": "code_review_classifier.v1",
        "counts": {k: len(v) for k, v in buckets.items()},
        "buckets": buckets,
    }


@mcp.tool()
def test_gen_small(
    function_name: str,
    path: str,
    framework: str = "pytest",
    behavior_summary: str = "",
) -> dict[str, Any]:
    """Generate a minimal unit-test skeleton for a target function."""
    if framework not in {"pytest", "unittest"}:
        raise ValueError("framework must be one of: pytest, unittest")
    target = _resolve_repo_path(path)
    if not target.is_file():
        raise FileNotFoundError(path)
    module = str(target.relative_to(REPO_PATH)).replace("/", ".")
    if module.endswith(".py"):
        module = module[:-3]
    module = module.replace(".__init__", "")
    if framework == "pytest":
        code = (
            f"from {module} import {function_name}\n\n"
            f"def test_{function_name}_basic():\n"
            f"    # {behavior_summary or 'TODO: define expected behavior'}\n"
            f"    result = {function_name}(...)\n"
            f"    assert result is not None\n"
        )
    else:
        code = (
            "import unittest\n"
            f"from {module} import {function_name}\n\n"
            "class GeneratedTest(unittest.TestCase):\n"
            f"    def test_{function_name}_basic(self):\n"
            f"        # {behavior_summary or 'TODO: define expected behavior'}\n"
            f"        result = {function_name}(...)\n"
            "        self.assertIsNotNone(result)\n"
        )
    return {"schema": "test_gen_small.v1", "framework": framework, "test_code": code}


@mcp.tool()
def vision_ocr_parser(
    image_path: str,
    language: str = "eng",
    max_chars: int = 5000,
) -> dict[str, Any]:
    """Offline OCR parser for local images using pytesseract when available."""
    path = _resolve_repo_path(image_path)
    if not path.is_file():
        raise FileNotFoundError(image_path)
    if Image is None or pytesseract is None:
        raise RuntimeError("vision OCR dependencies missing (Pillow/pytesseract)")
    img = Image.open(path)
    text = pytesseract.image_to_string(img, lang=language)
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    return {
        "schema": "vision_ocr_parser.v1",
        "image_path": image_path,
        "line_count": len(lines),
        "text": _trim_text(text, max_chars=max_chars),
    }


@mcp.tool()
def image_interpret(
    image_path: str,
    mode: str = "caption",
    language: str = "eng",
    use_local_model: bool = False,
    max_chars: int = 2000,
    output_profile: str | None = None,
) -> dict[str, Any]:
    """Interpret an image with constrained offline modes: ocr, caption, classify, ui_parse."""
    if mode not in {"ocr", "caption", "classify", "ui_parse"}:
        raise ValueError("mode must be one of: ocr, caption, classify, ui_parse")
    path = _resolve_repo_path(image_path)
    if not path.is_file():
        raise FileNotFoundError(image_path)
    profile = _default_output_profile(output_profile)

    warnings: list[str] = []
    features = _image_basic_features(path)
    ocr_text = ""
    if Image is not None and pytesseract is not None:
        try:
            with Image.open(path) as img:
                ocr_text = pytesseract.image_to_string(img, lang=language)
        except Exception as exc:
            warnings.append(f"OCR failed: {exc}")
    elif mode == "ocr":
        raise RuntimeError("vision OCR dependencies missing (Pillow/pytesseract)")

    lines = [ln.strip() for ln in ocr_text.splitlines() if ln.strip()]
    word_count = len(re.findall(r"\b\w+\b", ocr_text))
    aspect = float(features.get("aspect_ratio") or 0.0)
    mean_luma = features.get("mean_luma")

    label = "photo_like"
    confidence = 0.55
    if word_count > 40:
        label = "document_scan"
        confidence = 0.82
    elif word_count > 12 and 1.5 <= aspect <= 2.2:
        label = "ui_screenshot"
        confidence = 0.78
    elif word_count > 8:
        label = "diagram_or_slide"
        confidence = 0.68
    elif mean_luma is not None and float(mean_luma) > 210:
        label = "minimal_graphic"
        confidence = 0.62

    text_preview = " ".join(lines[:3]).strip()
    caption = (
        f"{label.replace('_', ' ')} image"
        f" ({int(features.get('width', 0))}x{int(features.get('height', 0))})."
    )
    if text_preview:
        caption += f" Visible text starts with: {text_preview[:180]}"

    ui_elements = []
    for idx, ln in enumerate(lines[:40], start=1):
        token_count = len(re.findall(r"\b\w+\b", ln))
        elem_type = "headline" if token_count <= 6 and len(ln) <= 60 else "text"
        ui_elements.append({"row": idx, "type": elem_type, "text": ln[:160]})

    summary = ""
    model_backend = ""
    if use_local_model and (caption or text_preview):
        prompt = (
            "Interpret this image using concise structured bullets:\n"
            f"- mode: {mode}\n"
            f"- heuristic_label: {label}\n"
            f"- caption: {caption}\n"
            f"- ocr_preview: {text_preview}\n"
            "- Return objective, key info, and uncertainty."
        )
        try:
            inferred = local_infer(
                prompt=prompt,
                task="image_interpret",
                backend="auto",
                output_profile="compact",
                max_tokens=220,
            )
            summary = str(inferred.get("output", "")).strip()
            model_backend = str(inferred.get("backend", ""))
        except Exception as exc:
            warnings.append(f"local model interpret failed: {exc}")
    if not summary:
        summary = caption

    payload = {
        "schema": "image_interpret.v1",
        "image_path": image_path,
        "mode": mode,
        "label": label,
        "confidence": round(confidence, 3),
        "features": features,
        "line_count": len(lines),
        "word_count": word_count,
        "warnings": warnings,
        "used_local_model": bool(model_backend),
        "model_backend": model_backend,
        "summary": _trim_text(summary, max_chars=max_chars),
        "ocr_text": _trim_text(ocr_text, max_chars=max_chars),
        "ui_elements": ui_elements,
    }
    if mode == "ocr":
        payload["summary"] = _trim_text(ocr_text, max_chars=max_chars)
    if mode == "classify":
        payload["summary"] = f"{label} ({round(confidence, 3)})"
    if mode == "ui_parse" and not ui_elements and ocr_text:
        payload["ui_elements"] = [{"row": 1, "type": "text", "text": _trim_text(ocr_text, max_chars=160)}]

    if profile == "compact":
        return {
            "schema": "image_interpret.compact.v1",
            "image_path": image_path,
            "mode": mode,
            "label": payload["label"],
            "confidence": payload["confidence"],
            "line_count": payload["line_count"],
            "warnings": warnings,
            "used_local_model": payload["used_local_model"],
            "summary": payload["summary"],
            "ui_elements": payload["ui_elements"][:12],
        }
    return payload


@mcp.tool()
def translation_small(
    text: str,
    source_lang: str = "en",
    target_lang: str = "de",
    mode: str = "lexical",
) -> dict[str, Any]:
    """Offline small translation helper with lexical fallback."""
    if mode not in {"lexical", "local_infer"}:
        raise ValueError("mode must be one of: lexical, local_infer")
    if mode == "lexical":
        translated = _simple_translate(text, source_lang, target_lang)
        return {
            "schema": "translation_small.v1",
            "mode": mode,
            "source_lang": source_lang,
            "target_lang": target_lang,
            "translated": translated,
        }
    infer = local_infer(
        prompt=f"Translate from {source_lang} to {target_lang}: {text}",
        task="translation",
        backend="auto",
        output_profile="compact",
        max_tokens=256,
    )
    return {
        "schema": "translation_small.v1",
        "mode": mode,
        "source_lang": source_lang,
        "target_lang": target_lang,
        "translated": infer.get("output", ""),
        "backend": infer.get("backend"),
    }


@mcp.tool()
def diagram_from_code(
    path: str = ".",
    diagram_type: str = "flowchart",
    max_nodes: int = 60,
    include_call_edges: bool = False,
    output_profile: str | None = None,
) -> dict[str, Any]:
    """Generate Mermaid diagrams from repository dependency/call metadata."""
    if diagram_type not in {"flowchart", "class", "sequence"}:
        raise ValueError("diagram_type must be one of: flowchart, class, sequence")
    if max_nodes < 1:
        raise ValueError("max_nodes must be >= 1")
    profile = _default_output_profile(output_profile)
    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)

    dep = dependency_map(
        path=path,
        recursive=True,
        output_profile="normal",
        fields=["from", "to"],
        limit=max_nodes * 4,
    )
    edges = dep.get("edges", []) if isinstance(dep, dict) else []
    nodes: set[str] = set()
    lines: list[str] = []

    if diagram_type == "flowchart":
        lines.append("flowchart LR")
        for e in edges:
            src = str(e.get("from", ""))
            dst = str(e.get("to", ""))
            if not src or not dst:
                continue
            nodes.add(src)
            nodes.add(dst)
            if len(nodes) > max_nodes:
                break
            sid = _mermaid_sanitize_id(src)
            did = _mermaid_sanitize_id(dst)
            lines.append(f'    {sid}["{src}"] --> {did}["{dst}"]')
    elif diagram_type == "class":
        lines.append("classDiagram")
        for e in edges:
            src = Path(str(e.get("from", ""))).stem
            dst = Path(str(e.get("to", ""))).stem
            if not src or not dst:
                continue
            nodes.add(src)
            nodes.add(dst)
            if len(nodes) > max_nodes:
                break
            lines.append(f"    class {src}")
            lines.append(f"    class {dst}")
            lines.append(f"    {src} --> {dst}")
    else:
        lines.append("sequenceDiagram")
        lines.append("    autonumber")
        for e in edges[:max_nodes]:
            src = Path(str(e.get("from", ""))).stem or "A"
            dst = Path(str(e.get("to", ""))).stem or "B"
            lines.append(f"    {src}->>{dst}: imports")

    if include_call_edges:
        cg = call_graph(
            path=path,
            output_profile="compact",
            fields=["path", "caller", "callee"],
            limit=max_nodes,
        )
        call_edges = cg.get("edges", []) if isinstance(cg, dict) else []
    else:
        call_edges = []

    mermaid = "\n".join(lines)
    result = {
        "schema": "diagram_from_code.v1",
        "diagram_type": diagram_type,
        "path": str(root.relative_to(REPO_PATH)),
        "node_count": len(nodes),
        "edge_count": len(edges),
        "mermaid": mermaid,
        "call_edges": call_edges if profile == "verbose" else [],
    }
    if profile == "compact":
        return {
            "schema": "diagram_from_code.compact.v1",
            "diagram_type": diagram_type,
            "node_count": result["node_count"],
            "edge_count": result["edge_count"],
            "mermaid": result["mermaid"],
        }
    return result


@mcp.tool()
def mermaid_lint_fix(
    mermaid_text: str,
    auto_fix: bool = True,
) -> dict[str, Any]:
    """Lint and optionally fix common Mermaid syntax issues."""
    if not mermaid_text.strip():
        raise ValueError("mermaid_text must not be empty")
    text = mermaid_text.replace("\t", "    ")
    issues: list[str] = []
    fixed = text

    header_re = re.compile(r"^\s*(flowchart|graph|classDiagram|sequenceDiagram)\b", re.MULTILINE)
    if not header_re.search(fixed):
        issues.append("Missing diagram header; prepended 'flowchart LR'.")
        if auto_fix:
            fixed = "flowchart LR\n" + fixed

    if "->" in fixed and "-->" not in fixed:
        issues.append("Potential invalid arrow syntax '->'; replacing with '-->'.")
        if auto_fix:
            fixed = fixed.replace("->", "-->")

    if "```" in fixed:
        issues.append("Remove markdown fences from raw mermaid input.")
        if auto_fix:
            fixed = fixed.replace("```mermaid", "").replace("```", "").strip()

    valid = True
    if "flowchart" in fixed and "-->" not in fixed and "---" not in fixed:
        valid = False
        issues.append("Flowchart has no edges.")

    return {
        "schema": "mermaid_lint_fix.v1",
        "valid": bool(valid),
        "issue_count": len(issues),
        "issues": issues,
        "fixed_mermaid": fixed if auto_fix else mermaid_text,
    }


@mcp.tool()
def drawio_generator(
    mode: str = "generate",
    nodes: list[dict[str, Any]] | None = None,
    edges: list[dict[str, Any]] | None = None,
    drawio_xml: str = "",
) -> dict[str, Any]:
    """Generate simple draw.io XML from graph data or parse XML back to graph."""
    if mode not in {"generate", "parse"}:
        raise ValueError("mode must be one of: generate, parse")

    if mode == "generate":
        ns = "https://app.diagrams.net"
        diagram_id = "diagram-1"
        nlist = nodes or []
        elist = edges or []
        root = [
            f'<mxfile host="{ns}">',
            f'  <diagram id="{diagram_id}" name="Page-1">',
            "    <mxGraphModel><root>",
            '      <mxCell id="0"/>',
            '      <mxCell id="1" parent="0"/>',
        ]
        for i, n in enumerate(nlist, start=2):
            nid = str(n.get("id", f"n{i}"))
            label = html.escape(str(n.get("label", nid)))
            x = int(n.get("x", 40 + (i - 2) * 40))
            y = int(n.get("y", 40 + (i - 2) * 20))
            root.append(
                f'      <mxCell id="{nid}" value="{label}" style="rounded=1;whiteSpace=wrap;html=1;" vertex="1" parent="1">'
                f'<mxGeometry x="{x}" y="{y}" width="140" height="60" as="geometry"/></mxCell>'
            )
        for i, e in enumerate(elist, start=1):
            eid = f"e{i}"
            src = html.escape(str(e.get("source", "")))
            dst = html.escape(str(e.get("target", "")))
            root.append(
                f'      <mxCell id="{eid}" edge="1" parent="1" source="{src}" target="{dst}"><mxGeometry relative="1" as="geometry"/></mxCell>'
            )
        root.extend(["    </root></mxGraphModel>", "  </diagram>", "</mxfile>"])
        xml = "\n".join(root)
        return {
            "schema": "drawio_generator.v1",
            "mode": mode,
            "node_count": len(nlist),
            "edge_count": len(elist),
            "drawio_xml": xml,
        }

    if not drawio_xml.strip():
        raise ValueError("drawio_xml must not be empty for parse mode")
    try:
        tree = ET.fromstring(drawio_xml)
    except ET.ParseError as exc:
        raise ValueError(f"invalid drawio xml: {exc}") from exc

    parsed_nodes: list[dict[str, Any]] = []
    parsed_edges: list[dict[str, Any]] = []
    for cell in tree.iter("mxCell"):
        if cell.attrib.get("vertex") == "1":
            parsed_nodes.append(
                {"id": cell.attrib.get("id", ""), "label": cell.attrib.get("value", "")}
            )
        if cell.attrib.get("edge") == "1":
            parsed_edges.append(
                {
                    "id": cell.attrib.get("id", ""),
                    "source": cell.attrib.get("source", ""),
                    "target": cell.attrib.get("target", ""),
                }
            )
    return {
        "schema": "drawio_generator.v1",
        "mode": mode,
        "node_count": len(parsed_nodes),
        "edge_count": len(parsed_edges),
        "nodes": parsed_nodes,
        "edges": parsed_edges,
    }


@mcp.tool()
def diagram_sync_check(
    source_paths: list[str],
    diagram_path: str,
    mode: str = "check",
    marker: str = "diagram-fingerprint",
) -> dict[str, Any]:
    """Check/update diagram freshness metadata against source file fingerprints."""
    if not source_paths:
        raise ValueError("source_paths must not be empty")
    if mode not in {"check", "update"}:
        raise ValueError("mode must be one of: check, update")
    diagram = _resolve_repo_path(diagram_path)
    if not diagram.is_file():
        raise FileNotFoundError(diagram_path)
    for p in source_paths:
        rp = _resolve_repo_path(p)
        if not rp.is_file():
            raise FileNotFoundError(p)

    fingerprint = _diagram_fingerprint(source_paths)
    text = diagram.read_text(encoding="utf-8", errors="replace")
    rx = re.compile(rf"{re.escape(marker)}:\s*([a-f0-9]{{64}})")
    m = rx.search(text)
    existing = m.group(1) if m else ""
    in_sync = existing == fingerprint

    if mode == "update":
        _require_mutations()
        new_line = f"{marker}: {fingerprint}"
        if m:
            text = rx.sub(new_line, text, count=1)
        else:
            text = text.rstrip() + "\n\n" + new_line + "\n"
        diagram.write_text(text, encoding="utf-8")
        in_sync = True

    return {
        "schema": "diagram_sync_check.v1",
        "mode": mode,
        "diagram_path": str(diagram.relative_to(REPO_PATH)),
        "source_paths": sorted(source_paths),
        "marker": marker,
        "fingerprint": fingerprint,
        "existing_fingerprint": existing,
        "in_sync": in_sync,
        "needs_update": not in_sync,
    }


@mcp.tool()
def local_model_status() -> dict[str, Any]:
    """Report local model configuration and endpoint availability."""
    status: dict[str, Any] = {
        "schema": "local_model_status.v1",
        "models_dir": str(LOCAL_MODELS_DIR),
        "models_dir_exists": _resolve_repo_path(".").exists() if str(LOCAL_MODELS_DIR).startswith(str(REPO_PATH)) else LOCAL_MODELS_DIR.exists(),
        "embed": {
            "backend": LOCAL_EMBED_BACKEND,
            "model": LOCAL_EMBED_MODEL,
            "dim": LOCAL_EMBED_DIM,
        },
        "infer": {
            "backend": LOCAL_INFER_BACKEND,
            "model": LOCAL_INFER_MODEL,
            "endpoint": LOCAL_INFER_ENDPOINT,
        },
    }
    if LOCAL_INFER_BACKEND == "endpoint":
        try:
            req = urllib.request.Request(
                LOCAL_INFER_ENDPOINT.replace("/api/generate", "/api/tags"),
                method="GET",
            )
            with _urlopen_with_host_certs(req, timeout=3) as resp:
                status["infer"]["endpoint_reachable"] = True
                status["infer"]["endpoint_status"] = getattr(resp, "status", 200)
        except Exception as exc:
            status["infer"]["endpoint_reachable"] = False
            status["infer"]["endpoint_error"] = str(exc)
    return status


@mcp.tool()
def local_embed(
    texts: list[str],
    backend: str = "auto",
    normalize: bool = True,
    output_profile: str | None = None,
    offset: int = 0,
    limit: int | None = None,
    compress: bool = False,
    store_result: bool = False,
) -> dict[str, Any]:
    """Create local offline embeddings for small specialized tasks."""
    if not texts:
        raise ValueError("texts must not be empty")
    profile = _default_output_profile(output_profile)
    selected, vectors = _local_embed_vectors(texts, backend=backend, normalize=normalize)
    rows = [{"index": i, "text": t, "vector": vectors[i]} for i, t in enumerate(texts)]
    rows = _paginate(rows, offset=offset, limit=limit)
    result: dict[str, Any] = {
        "schema": "local_embed.v1",
        "backend": selected,
        "count": len(rows),
        "dim": len(rows[0]["vector"]) if rows else 0,
        "rows": rows,
    }
    if profile == "compact":
        result = {
            "schema": "local_embed.compact.v1",
            "backend": selected,
            "count": len(rows),
            "dim": len(rows[0]["vector"]) if rows else 0,
            "rows": [{"index": r["index"]} for r in rows],
        }
    if compress and isinstance(result.get("rows"), list):
        result["rows_compressed"] = _compress_table(result["rows"])
        result.pop("rows", None)
    if store_result:
        result["result_id"] = _result_store_put("local_embed", result)
    return result


@mcp.tool()
def local_infer(
    prompt: str,
    task: str = "general",
    backend: str = "auto",
    model: str = "",
    max_tokens: int = 256,
    temperature: float = 0.2,
    system: str = "",
    output_profile: str | None = None,
    store_result: bool = False,
) -> dict[str, Any]:
    """Run local offline inference via endpoint or deterministic fallback."""
    if not prompt.strip():
        raise ValueError("prompt must not be empty")
    if max_tokens < 1:
        raise ValueError("max_tokens must be >= 1")
    if temperature < 0:
        raise ValueError("temperature must be >= 0")
    profile = _default_output_profile(output_profile)
    selected = backend.strip().lower()
    if selected == "auto":
        selected = LOCAL_INFER_BACKEND or "endpoint"
    model_name = model or LOCAL_INFER_MODEL or "local-default"

    if selected == "endpoint":
        try:
            text = _local_infer_via_endpoint(
                prompt=prompt,
                model=model_name,
                max_tokens=max_tokens,
                temperature=temperature,
                system=system,
            )
        except Exception as exc:
            _failure_record(
                command=["local_infer", "endpoint"],
                stderr=str(exc),
                category="local_infer",
                suggestion="Ensure local inference endpoint is running and reachable.",
            )
            text = ""
            selected = "fallback"
    if selected in {"fallback", "rule", "hash"}:
        optimized = prompt_optimize(prompt=prompt, mode="coding")
        text = optimized["optimized_prompt"][:max_tokens * 6]
    result = {
        "schema": "local_infer.v1",
        "backend": selected,
        "model": model_name,
        "task": task,
        "output": _trim_text(text),
        "ok": bool(text),
    }
    if profile == "compact":
        result = {
            "schema": "local_infer.compact.v1",
            "backend": selected,
            "model": model_name,
            "ok": bool(text),
            "output": _trim_text(text, max_chars=1200),
        }
    if store_result:
        result["result_id"] = _result_store_put("local_infer", result)
    return result


@mcp.tool()
def local_rerank(
    query: str,
    candidates: list[dict[str, Any]],
    top_k: int = 20,
    backend: str = "auto",
    output_profile: str | None = None,
) -> dict[str, Any]:
    """Rerank candidate items locally using offline embeddings."""
    if not query.strip():
        raise ValueError("query must not be empty")
    if top_k < 1:
        raise ValueError("top_k must be >= 1")
    if not candidates:
        return {"schema": "local_rerank.v1", "count": 0, "results": []}
    profile = _default_output_profile(output_profile)

    texts = []
    for c in candidates:
        txt = " ".join(
            str(c.get(k, "")) for k in ("path", "name", "match", "lineText", "kind")
        ).strip()
        texts.append(txt)
    selected, vectors = _local_embed_vectors([query, *texts], backend=backend, normalize=True)
    qv = vectors[0]
    scored = []
    for idx, cand in enumerate(candidates):
        score = _vec_cosine(qv, vectors[idx + 1])
        row = dict(cand)
        row["local_score"] = score
        scored.append(row)
    scored.sort(key=lambda x: float(x.get("local_score", 0.0)), reverse=True)
    scored = scored[:top_k]
    if profile == "compact":
        scored = [
            {"path": r.get("path"), "kind": r.get("kind"), "local_score": r.get("local_score")}
            for r in scored
        ]
    return {
        "schema": "local_rerank.v1",
        "backend": selected,
        "count": len(scored),
        "results": scored,
    }


@mcp.tool()
def token_budget_guard(
    max_output_chars: int | None = None,
    default_output_profile: str | None = None,
    reset: bool = False,
) -> dict[str, Any]:
    """Set or read global output budget/profile defaults."""
    if reset:
        payload = {
            "max_output_chars": MAX_OUTPUT_CHARS,
            "default_output_profile": "compact",
            "updated_at": _now_iso(),
        }
        _json_file_save(TOKEN_BUDGET_FILE, payload)
        return payload

    current = _token_budget_load()
    changed = False
    if max_output_chars is not None:
        if max_output_chars < 1:
            raise ValueError("max_output_chars must be >= 1")
        current["max_output_chars"] = max_output_chars
        changed = True
    if default_output_profile is not None:
        current["default_output_profile"] = _validate_output_profile(default_output_profile)
        changed = True
    if changed:
        current["updated_at"] = _now_iso()
        _json_file_save(TOKEN_BUDGET_FILE, current)
    return current


@mcp.tool()
def cache_control(
    mode: str = "stats",
    tool: str | None = None,
) -> dict[str, Any]:
    """Inspect or clear server-side tool cache entries."""
    if mode not in {"stats", "clear", "clear_tool"}:
        raise ValueError("mode must be one of: stats, clear, clear_tool")
    if mode == "stats":
        return {"mode": mode, **_cache_stats()}
    if mode == "clear":
        return {"mode": mode, **_cache_clear(None)}
    if not tool:
        raise ValueError("tool is required for clear_tool mode")
    return {"mode": mode, **_cache_clear(tool)}


@mcp.tool()
def result_handle(
    mode: str = "fetch",
    result_id: str = "",
    tool: str = "",
    value: Any = None,
    offset: int = 0,
    limit: int | None = None,
    fields: list[str] | None = None,
) -> dict[str, Any]:
    """Store/fetch/list/clear result handles for large payload workflows."""
    if mode not in {"store", "fetch", "list", "clear"}:
        raise ValueError("mode must be one of: store, fetch, list, clear")
    if mode == "store":
        rid = _result_store_put(tool=tool or "manual", value=value)
        return {"mode": mode, "result_id": rid}
    if mode == "list":
        payload = _result_store_load()
        items = []
        for rid, row in payload["results"].items():
            items.append(
                {
                    "result_id": rid,
                    "tool": row.get("tool"),
                    "created_at": row.get("created_at"),
                }
            )
        items = sorted(items, key=lambda x: str(x.get("created_at", "")), reverse=True)
        items = _paginate(items, offset=offset, limit=limit)
        items = _select_fields(items, fields)
        return {"mode": mode, "count": len(items), "results": items}
    if mode == "clear":
        _result_store_save({"results": {}})
        return {"mode": mode, "cleared": True}

    if not result_id:
        raise ValueError("result_id is required for fetch mode")
    row = _result_store_get(result_id)
    value_out = row.get("value")
    if isinstance(value_out, list):
        value_out = _paginate(value_out, offset=offset, limit=limit)
        if value_out and isinstance(value_out[0], dict):
            value_out = _select_fields(value_out, fields)
    return {
        "mode": mode,
        "result_id": result_id,
        "tool": row.get("tool"),
        "created_at": row.get("created_at"),
        "value": value_out,
    }


@mcp.tool()
def tool_benchmark(
    tools: list[str] | None = None,
    iterations: int = 3,
    warmup: int = 1,
) -> dict[str, Any]:
    """Benchmark representative tool invocations for latency and payload size."""
    if iterations < 1:
        raise ValueError("iterations must be >= 1")
    if warmup < 0:
        raise ValueError("warmup must be >= 0")

    catalog = {
        "find_paths": lambda: find_paths(path=".", recursive=True, max_entries=500, output_profile="compact"),
        "grep": lambda: grep(pattern="def ", path=".", recursive=True, max_matches=100, output_profile="compact"),
        "symbol_index": lambda: symbol_index(path=".", recursive=True, max_symbols=1000, output_profile="compact"),
        "dependency_map": lambda: dependency_map(path=".", recursive=True, max_files=1000, output_profile="compact", summary_mode="quick"),
        "call_graph": lambda: call_graph(path=".", recursive=True, max_edges=1000, output_profile="compact", summary_mode="quick"),
        "semantic_find": lambda: semantic_find(query="tool cache", path=".", max_results=20, output_profile="compact"),
        "tree_sitter_core": lambda: tree_sitter_core(path=".", mode="parse", max_files=20, max_nodes=500, output_profile="compact", summary_mode="quick"),
    }
    selected = tools or list(catalog.keys())
    unknown = [t for t in selected if t not in catalog]
    if unknown:
        raise ValueError(f"unknown benchmark tools: {', '.join(unknown)}")

    results: list[dict[str, Any]] = []
    for tool in selected:
        fn = catalog[tool]
        for _ in range(warmup):
            fn()
        latencies_ms: list[float] = []
        size_bytes: list[int] = []
        for _ in range(iterations):
            t0 = time.perf_counter()
            out = fn()
            t1 = time.perf_counter()
            latencies_ms.append((t1 - t0) * 1000.0)
            size_bytes.append(_payload_size_bytes(out))
        results.append(
            {
                "tool": tool,
                "iterations": iterations,
                "latency_ms_avg": round(sum(latencies_ms) / len(latencies_ms), 2),
                "latency_ms_p95": round(sorted(latencies_ms)[int(max(0, len(latencies_ms) * 0.95 - 1))], 2),
                "payload_bytes_avg": int(sum(size_bytes) / len(size_bytes)),
                "payload_bytes_max": int(max(size_bytes)),
            }
        )

    return {"schema": "tool_benchmark.v1", "results": results}


@mcp.tool()
def self_test(
    runner: str = "unittest",
    target: str = "tests",
    verbose: bool = True,
    timeout_seconds: int = 600,
    fail_fast: bool = False,
) -> dict[str, Any]:
    """Execute repository self-tests using unittest or pytest."""
    if runner not in {"unittest", "pytest"}:
        raise ValueError("runner must be one of: unittest, pytest")
    if timeout_seconds < 1:
        raise ValueError("timeout_seconds must be >= 1")

    out_cap = _token_budget_apply_max(None)
    _resolve_repo_path(target)

    if runner == "unittest":
        cmd = [sys.executable, "-m", "unittest"]
        target_path = _resolve_repo_path(target)
        if target_path.is_file() and target_path.suffix == ".py":
            rel_parent = str(target_path.parent.relative_to(REPO_PATH))
            cmd.extend(
                [
                    "discover",
                    "-s",
                    rel_parent if rel_parent else ".",
                    "-p",
                    target_path.name,
                ]
            )
            if verbose:
                cmd.append("-v")
            if fail_fast:
                cmd.append("-f")
        else:
            if verbose:
                cmd.append("-v")
            if fail_fast:
                cmd.append("-f")
            cmd.append(target)
    else:
        cmd = ["pytest"]
        if verbose:
            cmd.append("-v")
        else:
            cmd.append("-q")
        if fail_fast:
            cmd.append("-x")
        cmd.append(target)

    try:
        proc = subprocess.run(
            cmd,
            cwd=str(REPO_PATH),
            check=False,
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
        )
    except subprocess.TimeoutExpired as exc:
        _failure_record(
            command=cmd,
            stderr="self_test timed out",
            stdout=(exc.stdout or "") if isinstance(exc.stdout, str) else "",
            category="self_test",
            suggestion="Increase timeout_seconds or narrow the test target.",
        )
        return {
            "schema": "self_test.v1",
            "runner": runner,
            "target": target,
            "ok": False,
            "timeout": True,
            "exit_code": None,
            "command": cmd,
            "stdout": _trim_text((exc.stdout or "") if isinstance(exc.stdout, str) else "", max_chars=out_cap),
            "stderr": _trim_text((exc.stderr or "") if isinstance(exc.stderr, str) else "", max_chars=out_cap),
        }
    except FileNotFoundError as exc:
        _failure_record(
            command=cmd,
            stderr=str(exc),
            category="self_test",
            suggestion="Install the selected test runner in the runtime.",
        )
        return {
            "schema": "self_test.v1",
            "runner": runner,
            "target": target,
            "ok": False,
            "timeout": False,
            "exit_code": None,
            "command": cmd,
            "stdout": "",
            "stderr": str(exc),
        }

    if proc.returncode != 0:
        _failure_record(
            command=cmd,
            stderr=proc.stderr,
            stdout=proc.stdout,
            category="self_test",
            suggestion="Inspect failures and rerun with fail_fast=true for faster iteration.",
        )
    return {
        "schema": "self_test.v1",
        "runner": runner,
        "target": target,
        "ok": proc.returncode == 0,
        "timeout": False,
        "exit_code": proc.returncode,
        "command": cmd,
        "stdout": _trim_text(proc.stdout, max_chars=out_cap),
        "stderr": _trim_text(proc.stderr, max_chars=out_cap),
    }


@mcp.tool()
def output_size_guard(
    mode: str = "check",
    tools: list[str] | None = None,
    tolerance_ratio: float = 1.2,
    baseline_path: str = str(OUTPUT_BASELINE_FILE),
) -> dict[str, Any]:
    """Write/check baseline payload sizes to catch output-size regressions."""
    if mode not in {"write", "check"}:
        raise ValueError("mode must be one of: write, check")
    if tolerance_ratio < 1.0:
        raise ValueError("tolerance_ratio must be >= 1.0")

    bench = tool_benchmark(tools=tools, iterations=1, warmup=0)
    current = {r["tool"]: int(r["payload_bytes_max"]) for r in bench["results"]}
    baseline_file = _resolve_repo_path(baseline_path)

    if mode == "write":
        _require_mutations()
        baseline_file.parent.mkdir(parents=True, exist_ok=True)
        baseline_file.write_text(
            json.dumps({"generated_at": _now_iso(), "sizes": current}, indent=2),
            encoding="utf-8",
        )
        return {"mode": mode, "baseline_path": baseline_path, "sizes": current}

    if not baseline_file.is_file():
        raise FileNotFoundError(baseline_path)
    baseline = json.loads(baseline_file.read_text(encoding="utf-8"))
    prev = baseline.get("sizes", {})
    regressions: list[dict[str, Any]] = []
    for tool, cur in current.items():
        old = int(prev.get(tool, cur))
        if cur > int(old * tolerance_ratio):
            regressions.append({"tool": tool, "baseline": old, "current": cur})
    return {
        "mode": mode,
        "baseline_path": baseline_path,
        "ok": not regressions,
        "regressions": regressions,
        "current_sizes": current,
    }


@mcp.tool()
def failure_memory(
    mode: str = "get",
    category: str | None = None,
    contains: str | None = None,
    max_entries: int = 100,
    error_text: str = "",
    max_suggestions: int = 5,
) -> dict[str, Any]:
    """Unified failure memory access (`get` or `suggest`)."""
    if mode not in {"get", "suggest"}:
        raise ValueError("mode must be one of: get, suggest")
    if mode == "get":
        if max_entries < 1:
            raise ValueError("max_entries must be >= 1")
        payload = _failure_memory_load()
        entries_out: list[dict[str, Any]] = []
        needle = (contains or "").lower().strip()
        for entry in reversed(payload["entries"]):
            if category and entry.get("category") != category:
                continue
            if needle:
                hay = f"{entry.get('stderr', '')}\n{entry.get('stdout', '')}".lower()
                if needle not in hay:
                    continue
            entries_out.append(entry)
            if len(entries_out) >= max_entries:
                break
        return {"mode": mode, "count": len(entries_out), "entries": entries_out}

    if max_suggestions < 1:
        raise ValueError("max_suggestions must be >= 1")
    needle_terms = [t for t in re.split(r"\W+", error_text.lower()) if len(t) > 2]
    payload = _failure_memory_load()
    scored: list[tuple[int, dict[str, Any]]] = []
    for entry in payload["entries"]:
        hay = f"{entry.get('stderr', '')}\n{entry.get('stdout', '')}".lower()
        score = sum(1 for t in needle_terms if t in hay)
        if score <= 0:
            continue
        scored.append((score, entry))
    scored.sort(key=lambda x: x[0], reverse=True)
    suggestions: list[dict[str, Any]] = []
    for score, entry in scored[:max_suggestions]:
        suggestions.append(
            {
                "score": score,
                "category": entry.get("category"),
                "command": entry.get("command"),
                "suggestion": entry.get("suggestion"),
                "stderr": entry.get("stderr"),
            }
        )
    return {"mode": mode, "count": len(suggestions), "suggestions": suggestions}


@mcp.tool()
def edit_transaction(
    mode: str = "begin",
    transaction_id: str = "",
    label: str = "",
    changes: list[dict[str, Any]] | None = None,
    create_dirs: bool = True,
    delete_metadata: bool = False,
) -> dict[str, Any]:
    """Unified transaction tool (`begin|apply|validate|rollback|commit`)."""
    if mode not in {"begin", "apply", "validate", "rollback", "commit"}:
        raise ValueError("mode must be one of: begin, apply, validate, rollback, commit")

    if mode == "begin":
        _require_mutations()
        txn_id = uuid.uuid4().hex[:12]
        payload = {
            "id": txn_id,
            "label": label,
            "status": "open",
            "created_at": _now_iso(),
            "updated_at": _now_iso(),
            "changes": [],
            "backups": {},
        }
        _tx_save(txn_id, payload)
        return {"mode": mode, "transaction_id": txn_id, "status": "open"}

    if not transaction_id.strip():
        raise ValueError("transaction_id is required for this mode")

    if mode == "apply":
        _require_mutations()
        batch = changes or []
        if not batch:
            raise ValueError("changes must not be empty")
        tx = _tx_load(transaction_id)
        if tx.get("status") != "open":
            raise ValueError("transaction is not open")
        applied: list[str] = []
        for change in batch:
            path = change.get("path")
            content = change.get("content")
            if not isinstance(path, str) or not isinstance(content, str):
                raise ValueError("each change requires string path and content")
            file_path = _resolve_repo_path(path)
            rel = str(file_path.relative_to(REPO_PATH))

            if rel not in tx["backups"]:
                if file_path.exists() and file_path.is_file():
                    tx["backups"][rel] = {
                        "existed": True,
                        "content": file_path.read_text(encoding="utf-8", errors="replace"),
                    }
                else:
                    tx["backups"][rel] = {"existed": False, "content": ""}

            if create_dirs:
                file_path.parent.mkdir(parents=True, exist_ok=True)
            file_path.write_text(content, encoding="utf-8")
            tx["changes"].append({"path": rel, "bytes": len(content.encode("utf-8"))})
            applied.append(rel)

        tx["updated_at"] = _now_iso()
        _tx_save(transaction_id, tx)
        return {
            "mode": mode,
            "transaction_id": transaction_id,
            "applied": applied,
            "change_count": len(applied),
        }

    if mode == "validate":
        tx = _tx_load(transaction_id)
        changed_paths = sorted({c["path"] for c in tx.get("changes", [])})
        py_files = [p for p in changed_paths if p.endswith(".py")]
        compile_errors: list[dict[str, Any]] = []
        for rel in py_files[:200]:
            file_path = _resolve_repo_path(rel)
            proc = subprocess.run(
                [sys.executable, "-m", "py_compile", str(file_path)],
                check=False,
                capture_output=True,
                text=True,
            )
            if proc.returncode != 0:
                compile_errors.append({"path": rel, "stderr": _trim_text(proc.stderr)})
        return {
            "mode": mode,
            "transaction_id": transaction_id,
            "status": tx.get("status"),
            "changed_paths": changed_paths,
            "python_files_checked": len(py_files[:200]),
            "compile_error_count": len(compile_errors),
            "compile_errors": compile_errors,
        }

    if mode == "rollback":
        _require_mutations()
        tx = _tx_load(transaction_id)
        backups = tx.get("backups", {})
        restored: list[str] = []
        for rel, backup in backups.items():
            file_path = _resolve_repo_path(rel)
            existed = bool(backup.get("existed"))
            if existed:
                file_path.parent.mkdir(parents=True, exist_ok=True)
                file_path.write_text(str(backup.get("content", "")), encoding="utf-8")
            else:
                if file_path.exists():
                    file_path.unlink()
            restored.append(rel)
        tx["status"] = "rolled_back"
        tx["updated_at"] = _now_iso()
        _tx_save(transaction_id, tx)
        return {
            "mode": mode,
            "transaction_id": transaction_id,
            "status": "rolled_back",
            "restored": restored,
        }

    tx = _tx_load(transaction_id)
    tx["status"] = "committed"
    tx["updated_at"] = _now_iso()
    _tx_save(transaction_id, tx)
    if delete_metadata:
        _tx_path(transaction_id).unlink(missing_ok=True)
    return {
        "mode": mode,
        "transaction_id": transaction_id,
        "status": "committed",
        "metadata_deleted": delete_metadata,
    }


@mcp.tool()
def impact_tests(
    base_ref: str = "HEAD~1",
    head_ref: str = "HEAD",
    max_tests: int = 300,
    output_profile: str | None = None,
) -> dict[str, Any]:
    """Select impacted tests using changed files and dependency edges."""
    _require_git_repo()
    if max_tests < 1:
        raise ValueError("max_tests must be >= 1")
    profile = _default_output_profile(output_profile)
    diff_out = _git("diff", "--name-only", f"{base_ref}...{head_ref}").stdout.strip()
    changed = [line.strip() for line in diff_out.splitlines() if line.strip()]

    dep = dependency_map(path=".", recursive=True, include_stdlib=False, output_profile="normal")
    reverse_edges: dict[str, set[str]] = {}
    for edge in dep.get("edges", []):
        reverse_edges.setdefault(edge["to"], set()).add(edge["from"])

    impacted: set[str] = set(changed)
    queue: list[str] = list(changed)
    while queue:
        cur = queue.pop(0)
        for src in reverse_edges.get(cur, set()):
            if src not in impacted:
                impacted.add(src)
                queue.append(src)

    tests: list[str] = []
    for rel in sorted(impacted):
        p = Path(rel)
        if "test" in p.name.lower() and p.suffix == ".py":
            tests.append(rel)
            continue
        if p.suffix == ".py":
            for cand in (f"tests/test_{p.stem}.py", f"tests/{p.stem}_test.py"):
                resolved = _resolve_repo_path(cand)
                if resolved.is_file():
                    tests.append(cand)
    deduped: list[str] = []
    seen: set[str] = set()
    for t in tests:
        if t in seen:
            continue
        seen.add(t)
        deduped.append(t)
        if len(deduped) >= max_tests:
            break

    result = {
        "base_ref": base_ref,
        "head_ref": head_ref,
        "changed_files": changed,
        "impacted_files": sorted(impacted),
        "tests": deduped,
    }
    if profile == "compact":
        return {"test_count": len(deduped), "tests": deduped}
    return result


@mcp.tool()
def api_surface_snapshot(
    path: str = ".",
    snapshot_path: str = str(API_SNAPSHOT_FILE),
    mode: str = "write",
    include_private: bool = False,
) -> dict[str, Any]:
    """Write or check public Python API surface snapshots."""
    _require_git_repo()
    if mode not in {"write", "check"}:
        raise ValueError("mode must be one of: write, check")

    symbols = symbol_index(
        path=path,
        include_private=include_private,
        recursive=True,
        max_symbols=20000,
        output_profile="normal",
    )
    public_symbols = [
        {
            "path": s["path"],
            "name": s["name"],
            "kind": s["kind"],
        }
        for s in symbols
        if include_private or not str(s["name"]).startswith("_")
    ]
    public_symbols.sort(key=lambda x: (x["path"], x["name"], x["kind"]))
    snap_file = _resolve_repo_path(snapshot_path)

    if mode == "write":
        _require_mutations()
        snap_file.parent.mkdir(parents=True, exist_ok=True)
        snap_file.write_text(
            json.dumps({"generated_at": _now_iso(), "symbols": public_symbols}, indent=2),
            encoding="utf-8",
        )
        return {"mode": "write", "snapshot_path": snapshot_path, "symbol_count": len(public_symbols)}

    if not snap_file.is_file():
        raise FileNotFoundError(snapshot_path)
    baseline = json.loads(snap_file.read_text(encoding="utf-8"))
    baseline_symbols = baseline.get("symbols", [])
    base_set = {(x["path"], x["name"], x["kind"]) for x in baseline_symbols if isinstance(x, dict)}
    cur_set = {(x["path"], x["name"], x["kind"]) for x in public_symbols}
    removed = sorted(base_set - cur_set)
    added = sorted(cur_set - base_set)
    return {
        "mode": "check",
        "snapshot_path": snapshot_path,
        "removed_count": len(removed),
        "added_count": len(added),
        "removed": [{"path": p, "name": n, "kind": k} for (p, n, k) in removed],
        "added": [{"path": p, "name": n, "kind": k} for (p, n, k) in added],
    }


@mcp.tool()
def workspace_facts(refresh: bool = True) -> dict[str, Any]:
    """Get or refresh lightweight workspace facts."""
    facts_path = Path(".build/memory/workspace_facts.json")
    if not refresh:
        payload = _json_file_load(facts_path, {})
        if payload:
            return payload

    files = find_paths(path=".", recursive=True, file_type="file", max_entries=10000, output_profile="compact")
    ext_counts: dict[str, int] = {}
    for rel in files:
        ext = Path(rel).suffix.lower() or "<none>"
        ext_counts[ext] = ext_counts.get(ext, 0) + 1
    top_ext = sorted(
        [{"extension": k, "count": v} for k, v in ext_counts.items()],
        key=lambda x: x["count"],
        reverse=True,
    )[:20]
    payload = {
        "generated_at": _now_iso(),
        "is_git_repo": _is_git_repo(),
        "file_count": len(files),
        "top_extensions": top_ext,
        "has_tests_dir": _resolve_repo_path("tests").exists(),
        "has_readme": _resolve_repo_path("README.md").exists(),
        "default_output_profile": _token_budget_load()["default_output_profile"],
    }
    _json_file_save(facts_path, payload)
    return payload


@mcp.tool()
def risk_scoring(
    ref: str | None = None,
    staged: bool = False,
    pathspec: str | None = None,
) -> dict[str, Any]:
    """Score change risk using path and churn heuristics."""
    summary = summarize_diff(ref=ref, staged=staged, pathspec=pathspec, output_profile="normal")
    score = 0
    reasons: list[str] = []
    file_count = int(summary["file_count"])
    add = int(summary["total_added"])
    delete = int(summary["total_deleted"])
    churn = add + delete

    if file_count > 20:
        score += 2
        reasons.append("large file count")
    if churn > 500:
        score += 2
        reasons.append("high churn")
    if churn > 1500:
        score += 2
        reasons.append("very high churn")

    risky_files = summary.get("risk_flags", {}).get("risky_files", [])
    if risky_files:
        score += min(4, len(risky_files))
        reasons.append("sensitive file changes")

    todo_adds = int(summary.get("risk_flags", {}).get("todo_like_additions", 0))
    if todo_adds > 0:
        score += 1
        reasons.append("todo/fixme additions")

    level = "low"
    if score >= 6:
        level = "high"
    elif score >= 3:
        level = "medium"
    return {
        "risk_score": score,
        "risk_level": level,
        "reasons": reasons,
        "summary": summary,
    }


@mcp.tool()
def doc_sync_check(
    base_ref: str = "HEAD~1",
    head_ref: str = "HEAD",
    doc_globs: list[str] | None = None,
) -> dict[str, Any]:
    """Check whether docs changed when code/API-like files changed."""
    _require_git_repo()
    docs = doc_globs or ["README.md", "docs/**", "**/*.md"]
    diff_out = _git("diff", "--name-only", f"{base_ref}...{head_ref}").stdout.strip()
    changed = [line.strip() for line in diff_out.splitlines() if line.strip()]
    doc_changed: list[str] = []
    code_changed: list[str] = []
    for rel in changed:
        if any(fnmatch.fnmatch(rel, g) for g in docs):
            doc_changed.append(rel)
            continue
        if rel.endswith((".py", ".ts", ".tsx", ".js", ".go", ".rs")):
            code_changed.append(rel)
    needs_docs = bool(code_changed) and not bool(doc_changed)
    suggestions: list[str] = []
    if needs_docs:
        suggestions.append("Update README.md with behavioral/API changes.")
        if _resolve_repo_path("docs").exists():
            suggestions.append("Add or update docs under docs/.")
    return {
        "base_ref": base_ref,
        "head_ref": head_ref,
        "changed_count": len(changed),
        "code_changed": code_changed,
        "doc_changed": doc_changed,
        "needs_docs_update": needs_docs,
        "suggestions": suggestions,
    }


@mcp.tool()
def tree_sitter_core(
    path: str = ".",
    mode: str = "status",
    language: str = "auto",
    node_types: list[str] | None = None,
    text_pattern: str | None = None,
    recursive: bool = True,
    max_files: int = 200,
    max_nodes: int = 5000,
    output_profile: str | None = None,
    fields: list[str] | None = None,
    offset: int = 0,
    limit: int | None = None,
    adaptive_limits: bool = True,
    summary_mode: str = "full",
    compress: bool = False,
    store_result: bool = False,
) -> dict[str, Any]:
    """Parse/search syntax trees via Tree-sitter when available."""
    if mode not in {"status", "parse", "search"}:
        raise ValueError("mode must be one of: status, parse, search")
    if max_files < 1 or max_nodes < 1:
        raise ValueError("max_files and max_nodes must be >= 1")
    if summary_mode not in {"full", "quick"}:
        raise ValueError("summary_mode must be one of: full, quick")
    if adaptive_limits:
        max_files = _adaptive_limit(max_files, soft_cap=120)
        max_nodes = _adaptive_limit(max_nodes, soft_cap=1500)
    profile = _default_output_profile(output_profile)
    available = _tree_sitter_available()
    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)

    if mode == "status":
        return {"available": available, "engine": "tree_sitter_languages"}

    fingerprint = _fingerprint_path(root, recursive=recursive, max_files=3000)
    cache_key = json.dumps(
        {
            "path": str(root.relative_to(REPO_PATH)),
            "mode": mode,
            "language": language,
            "node_types": node_types or [],
            "text_pattern": text_pattern or "",
            "recursive": recursive,
            "max_files": max_files,
            "max_nodes": max_nodes,
            "fingerprint": fingerprint,
            "available": available,
        },
        sort_keys=True,
    )
    cached = _cache_get("tree_sitter_core", cache_key)
    if isinstance(cached, dict):
        matched_files = int(cached.get("file_count", 0))
        total_nodes = int(cached.get("node_count", 0))
        files = list(cached.get("files", []))
    else:
        matched_files = 0
        total_nodes = 0
        files: list[dict[str, Any]] = []
        regex = re.compile(text_pattern, re.IGNORECASE) if text_pattern else None

        for candidate in _iter_candidate_files(root, recursive=recursive):
            ext = candidate.suffix.lower()
            detected = _tree_sitter_language_for_ext(ext)
            if not detected:
                continue
            lang = detected if language == "auto" else language
            if language != "auto" and lang != detected:
                continue
            source = candidate.read_text(encoding="utf-8", errors="replace")
            rel = str(candidate.relative_to(REPO_PATH)).replace("\\", "/")
            nodes: list[dict[str, Any]] = []

            if available:
                try:
                    nodes = _tree_sitter_parse_nodes(
                        source=source,
                        language=lang,
                        node_types=node_types,
                        max_nodes=max_nodes - total_nodes,
                    )
                except Exception:
                    nodes = []
            if not nodes and lang == "python":
                ast_hits = ast_search(
                    path=rel,
                    node_type=(node_types[0] if node_types else "FunctionDef"),
                    max_results=max(1, max_nodes - total_nodes),
                )
                nodes = [
                    {
                        "type": hit["node_type"],
                        "start_line": hit["line"],
                        "start_column": hit["column"],
                        "end_line": hit["end_line"],
                        "end_column": hit["column"],
                    }
                    for hit in ast_hits
                ]

            if regex:
                lines = source.splitlines()
                filtered: list[dict[str, Any]] = []
                for n in nodes:
                    s = max(1, int(n["start_line"]))
                    e = min(len(lines), int(n["end_line"]))
                    snippet = "\n".join(lines[s - 1 : e])
                    if regex.search(snippet):
                        filtered.append(n)
                nodes = filtered

            if not nodes:
                continue
            matched_files += 1
            total_nodes += len(nodes)
            files.append({"path": rel, "language": lang, "node_count": len(nodes), "nodes": nodes})
            if matched_files >= max_files or total_nodes >= max_nodes:
                break
        _cache_set(
            "tree_sitter_core",
            cache_key,
            {"file_count": matched_files, "node_count": total_nodes, "files": files},
        )

    files = _paginate(files, offset=offset, limit=limit)
    if profile == "compact" and not fields:
        files = [
            {"path": f.get("path"), "language": f.get("language"), "node_count": f.get("node_count")}
            for f in files
        ]
    else:
        files = _select_fields(files, fields)

    result = {
        "schema": "tree_sitter_core.v1",
        "available": available,
        "mode": mode,
        "path": str(root.relative_to(REPO_PATH)),
        "file_count": matched_files,
        "node_count": total_nodes,
        "files": files,
    }
    if summary_mode == "quick":
        result = {
            "schema": "tree_sitter_core.quick.v1",
            "available": available,
            "mode": mode,
            "path": str(root.relative_to(REPO_PATH)),
            "file_count": matched_files,
            "node_count": total_nodes,
        }
    if compress and isinstance(result.get("files"), list):
        result["files_compressed"] = _compress_table(result["files"])
        result.pop("files", None)
    if store_result:
        result["result_id"] = _result_store_put("tree_sitter_core", result)
    return result


@mcp.tool()
def repo_index_daemon(
    mode: str = "refresh",
    path: str = ".",
    query: str = "",
    recursive: bool = True,
    include_hashes: bool = False,
    max_files: int = 5000,
    output_profile: str | None = None,
    fields: list[str] | None = None,
    offset: int = 0,
    limit: int | None = None,
    adaptive_limits: bool = True,
    summary_mode: str = "full",
    compress: bool = False,
    store_result: bool = False,
    incremental: bool = True,
) -> dict[str, Any]:
    """Build/read/query a persistent repository index keyed by file metadata."""
    if mode not in {"refresh", "read", "query"}:
        raise ValueError("mode must be one of: refresh, read, query")
    if max_files < 1:
        raise ValueError("max_files must be >= 1")
    if summary_mode not in {"full", "quick"}:
        raise ValueError("summary_mode must be one of: full, quick")
    if adaptive_limits:
        max_files = _adaptive_limit(max_files, soft_cap=2500)
    profile = _default_output_profile(output_profile)
    index_path = _resolve_repo_path(str(REPO_INDEX_FILE))

    if mode in {"read", "query"}:
        if not index_path.is_file():
            raise FileNotFoundError(str(REPO_INDEX_FILE))
        index_payload = json.loads(index_path.read_text(encoding="utf-8"))
        if mode == "query":
            value = _query_value(index_payload, query) if query.strip() else index_payload
            if isinstance(value, list):
                value = _paginate(value, offset=offset, limit=limit)
                if value and isinstance(value[0], dict):
                    value = _select_fields(value, fields)
            return {
                "mode": mode,
                "query": query,
                "value_json": _trim_text(json.dumps(value, indent=2, ensure_ascii=True)),
                "value": value if profile != "compact" else None,
            }
        if isinstance(index_payload.get("files"), list):
            files = _paginate(index_payload["files"], offset=offset, limit=limit)
            if files and isinstance(files[0], dict):
                files = _select_fields(files, fields)
            index_payload["files"] = files
        if profile == "compact":
            compact = {
                "schema": "repo_index_daemon.compact.v1",
                "mode": mode,
                "generated_at": index_payload.get("generated_at"),
                "file_count": index_payload.get("file_count", 0),
                "symbol_count": index_payload.get("symbol_count", 0),
                "dependency_edge_count": index_payload.get("dependency_edge_count", 0),
            }
            if store_result:
                compact["result_id"] = _result_store_put("repo_index_daemon", compact)
            return compact
        if summary_mode == "quick":
            quick = {
                "schema": "repo_index_daemon.quick.v1",
                "mode": mode,
                "generated_at": index_payload.get("generated_at"),
                "file_count": index_payload.get("file_count", 0),
                "symbol_count": index_payload.get("symbol_count", 0),
            }
            if store_result:
                quick["result_id"] = _result_store_put("repo_index_daemon", quick)
            return quick
        if compress and isinstance(index_payload.get("files"), list):
            index_payload["files_compressed"] = _compress_table(index_payload["files"])
            index_payload.pop("files", None)
        if store_result:
            index_payload["result_id"] = _result_store_put("repo_index_daemon", index_payload)
        return index_payload

    root = _resolve_repo_path(path)
    if not root.exists():
        raise FileNotFoundError(path)

    existing_payload = None
    if incremental and index_path.is_file():
        try:
            existing_payload = json.loads(index_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            existing_payload = None

    files_meta: list[dict[str, Any]] = []
    for candidate in _iter_candidate_files(root, recursive=recursive):
        rel = str(candidate.relative_to(REPO_PATH)).replace("\\", "/")
        stat = candidate.stat()
        entry = {
            "path": rel,
            "size": int(stat.st_size),
            "mtime_ns": int(stat.st_mtime_ns),
        }
        if include_hashes:
            try:
                entry["sha256"] = _file_sha256(candidate)
            except OSError:
                entry["sha256"] = ""
        files_meta.append(entry)
        if len(files_meta) >= max_files:
            break

    changed_paths: list[str] = []
    if existing_payload and isinstance(existing_payload.get("files"), list):
        prev = {
            f.get("path"): (f.get("size"), f.get("mtime_ns"))
            for f in existing_payload["files"]
            if isinstance(f, dict) and isinstance(f.get("path"), str)
        }
        for row in files_meta:
            p = row["path"]
            sig = (row.get("size"), row.get("mtime_ns"))
            if prev.get(p) != sig:
                changed_paths.append(p)

    reuse_prev_graphs = bool(existing_payload) and len([p for p in changed_paths if p.endswith(".py")]) == 0
    if reuse_prev_graphs:
        symbols = existing_payload.get("symbols", []) if isinstance(existing_payload.get("symbols"), list) else []
        dep_edges = existing_payload.get("dependencies", []) if isinstance(existing_payload.get("dependencies"), list) else []
        call_edges = existing_payload.get("call_edges", []) if isinstance(existing_payload.get("call_edges"), list) else []
        dep = {"edge_count": len(dep_edges), "edges": dep_edges}
        call = {"edge_count": len(call_edges), "edges": call_edges}
    else:
        symbols = symbol_index(
            path=path,
            include_private=False,
            recursive=recursive,
            max_symbols=20000,
            output_profile="compact",
        )
        dep = dependency_map(
            path=path,
            recursive=recursive,
            include_stdlib=False,
            max_files=max_files,
            output_profile="compact",
        )
        call = call_graph(
            path=path,
            recursive=recursive,
            max_edges=20000,
            output_profile="compact",
        )

    payload: dict[str, Any] = {
        "schema": "repo_index_daemon.v1",
        "generated_at": _now_iso(),
        "path": str(root.relative_to(REPO_PATH)),
        "file_count": len(files_meta),
        "symbol_count": (
            int(existing_payload.get("symbol_count", 0))
            if reuse_prev_graphs and existing_payload
            else len(symbols)
        ),
        "dependency_edge_count": int(dep.get("edge_count", 0)),
        "call_edge_count": int(call.get("edge_count", 0)),
        "files": files_meta if profile != "compact" else files_meta[:300],
        "symbols": symbols if profile == "verbose" else [],
        "dependencies": dep.get("edges", []) if profile == "verbose" else [],
        "call_edges": call.get("edges", []) if profile == "verbose" else [],
        "incremental": incremental,
        "changed_paths_count": len(changed_paths),
    }
    if _is_git_repo():
        payload["git_head"] = _git("rev-parse", "HEAD").stdout.strip()
        payload["git_branch"] = _git("branch", "--show-current").stdout.strip()

    index_path.parent.mkdir(parents=True, exist_ok=True)
    index_path.write_text(json.dumps(payload, indent=2, sort_keys=True), encoding="utf-8")
    result = {
        "schema": "repo_index_daemon.refresh.v1",
        "mode": mode,
        "index_path": str(REPO_INDEX_FILE),
        "generated_at": payload["generated_at"],
        "file_count": payload["file_count"],
        "symbol_count": payload["symbol_count"],
        "dependency_edge_count": payload["dependency_edge_count"],
        "call_edge_count": payload["call_edge_count"],
        "incremental": incremental,
        "changed_paths_count": len(changed_paths),
    }
    if summary_mode == "quick":
        result = {
            "schema": "repo_index_daemon.quick.v1",
            "mode": mode,
            "generated_at": payload["generated_at"],
            "file_count": payload["file_count"],
            "symbol_count": payload["symbol_count"],
            "incremental": incremental,
            "changed_paths_count": len(changed_paths),
        }
    if compress:
        result["files_compressed"] = _compress_table(files_meta[:500])
    if store_result:
        result["result_id"] = _result_store_put("repo_index_daemon", result)
    return result


@mcp.tool()
def self_check_pipeline(
    base_ref: str = "HEAD~1",
    head_ref: str = "HEAD",
    run_test_execution: bool = True,
    run_impact_tests: bool = True,
    run_doc_check: bool = True,
    run_api_check: bool = True,
    run_risk_check: bool = True,
    run_compile_check: bool = True,
    snapshot_path: str = str(API_SNAPSHOT_FILE),
    max_compile_files: int = 300,
    summary_mode: str = "quick",
) -> dict[str, Any]:
    """Run a single-call quality gate pipeline and return structured results."""
    _require_git_repo()
    if max_compile_files < 1:
        raise ValueError("max_compile_files must be >= 1")
    if summary_mode not in {"quick", "full"}:
        raise ValueError("summary_mode must be one of: quick, full")

    result: dict[str, Any] = {
        "base_ref": base_ref,
        "head_ref": head_ref,
        "started_at": _now_iso(),
        "checks": {},
        "ok": True,
    }

    diff_out = _git("diff", "--name-only", f"{base_ref}...{head_ref}").stdout.strip()
    changed = [line.strip() for line in diff_out.splitlines() if line.strip()]
    result["changed_files"] = changed

    if run_compile_check:
        compile_errors: list[dict[str, Any]] = []
        py_files = [p for p in changed if p.endswith(".py")][:max_compile_files]
        for rel in py_files:
            proc = subprocess.run(
                [sys.executable, "-m", "py_compile", str(_resolve_repo_path(rel))],
                check=False,
                capture_output=True,
                text=True,
            )
            if proc.returncode != 0:
                compile_errors.append({"path": rel, "stderr": _trim_text(proc.stderr)})
        result["checks"]["compile"] = {
            "checked_files": len(py_files),
            "error_count": len(compile_errors),
            "errors": compile_errors if summary_mode == "full" else compile_errors[:10],
        }
        if compile_errors:
            result["ok"] = False

    if run_risk_check:
        risk = risk_scoring(ref=f"{base_ref}...{head_ref}")
        result["checks"]["risk"] = risk
        if risk.get("risk_level") == "high":
            result["ok"] = False

    if run_doc_check:
        doc = doc_sync_check(base_ref=base_ref, head_ref=head_ref)
        result["checks"]["docs"] = doc
        if doc.get("needs_docs_update"):
            result["ok"] = False

    if run_api_check:
        if _resolve_repo_path(snapshot_path).is_file():
            api = api_surface_snapshot(
                mode="check",
                snapshot_path=snapshot_path,
                include_private=False,
            )
            result["checks"]["api"] = api
            if api.get("removed_count", 0) > 0:
                result["ok"] = False
        else:
            result["checks"]["api"] = {
                "skipped": True,
                "reason": f"snapshot missing: {snapshot_path}",
            }

    if run_impact_tests:
        impacts = impact_tests(
            base_ref=base_ref,
            head_ref=head_ref,
            output_profile="compact",
        )
        result["checks"]["impact_tests"] = impacts

    if run_test_execution:
        selected = impact_tests(
            base_ref=base_ref,
            head_ref=head_ref,
            output_profile="compact",
        ).get("tests", [])
        test_cmd = ["pytest", "-q", *selected] if selected else ["pytest", "-q"]
        exec_result = command_runner(
            command=test_cmd,
            cwd=".",
            timeout_seconds=300,
        )
        result["checks"]["test_execution"] = {
            "ok": exec_result.get("ok", False),
            "exit_code": exec_result.get("exit_code"),
            "command": exec_result.get("command"),
            "selected_tests": selected,
            "stderr": exec_result.get("stderr", ""),
            "timeout": exec_result.get("timeout", False),
        }
        if not exec_result.get("ok", False):
            result["ok"] = False

    result["finished_at"] = _now_iso()
    if summary_mode == "quick":
        return {
            "schema": "self_check_pipeline.quick.v1",
            "base_ref": base_ref,
            "head_ref": head_ref,
            "ok": result["ok"],
            "changed_file_count": len(changed),
            "checks": {
                name: {
                    k: v
                    for k, v in data.items()
                    if k in {"ok", "exit_code", "error_count", "needs_docs_update", "risk_level", "removed_count", "added_count", "timeout", "checked_files"}
                }
                for name, data in result["checks"].items()
                if isinstance(data, dict)
            },
        }
    return result


@mcp.tool()
def memory_upsert(
    namespace: str,
    key: str,
    value: Any,
    ttl_days: int | None = None,
    confidence: float = 1.0,
    source: str = "agent",
    tags: list[str] | None = None,
) -> dict[str, Any]:
    """Create or update a structured context memory record."""
    _require_mutations()
    if not namespace.strip() or not key.strip():
        raise ValueError("namespace and key must not be empty")
    if confidence < 0 or confidence > 1:
        raise ValueError("confidence must be in range [0, 1]")

    payload = _memory_load()
    entries = payload["entries"]
    now_iso = _now_iso()
    expires_at = _to_iso_expiry(ttl_days)
    updated = False
    for entry in entries:
        if entry.get("namespace") == namespace and entry.get("key") == key:
            entry["value"] = value
            entry["confidence"] = confidence
            entry["source"] = source
            entry["tags"] = tags or []
            entry["updated_at"] = now_iso
            entry["expires_at"] = expires_at
            updated = True
            break

    if not updated:
        entries.append(
            {
                "namespace": namespace,
                "key": key,
                "value": value,
                "confidence": confidence,
                "source": source,
                "tags": tags or [],
                "created_at": now_iso,
                "updated_at": now_iso,
                "expires_at": expires_at,
            }
        )

    _memory_save(payload)
    return {
        "path": str(MEMORY_FILE),
        "namespace": namespace,
        "key": key,
        "updated": True,
        "expires_at": expires_at,
    }


@mcp.tool()
def memory_get(
    namespace: str | None = None,
    key: str | None = None,
    include_expired: bool = False,
    max_entries: int = 200,
) -> dict[str, Any]:
    """Read context memory entries with namespace/key filters."""
    if max_entries < 1:
        raise ValueError("max_entries must be >= 1")
    payload = _memory_load()
    now = datetime.now(timezone.utc)
    entries_out: list[dict[str, Any]] = []

    for entry in payload["entries"]:
        if namespace is not None and entry.get("namespace") != namespace:
            continue
        if key is not None and entry.get("key") != key:
            continue
        expired = _is_expired(entry.get("expires_at"), now)
        if expired and not include_expired:
            continue
        copied = dict(entry)
        copied["expired"] = expired
        entries_out.append(copied)
        if len(entries_out) >= max_entries:
            break

    return {
        "path": str(MEMORY_FILE),
        "count": len(entries_out),
        "entries": entries_out,
    }


@mcp.tool()
def memory_validate(
    validate_paths: bool = True,
    drop_expired: bool = False,
    max_entries: int = 5000,
) -> dict[str, Any]:
    """Validate memory freshness and optionally prune expired records."""
    if max_entries < 1:
        raise ValueError("max_entries must be >= 1")
    payload = _memory_load()
    entries = payload["entries"][:max_entries]
    now = datetime.now(timezone.utc)

    stale: list[dict[str, Any]] = []
    kept: list[dict[str, Any]] = []
    dropped = 0

    for entry in entries:
        expired = _is_expired(entry.get("expires_at"), now)
        if expired and drop_expired:
            dropped += 1
            continue

        record = dict(entry)
        record["expired"] = expired
        record["stale_paths"] = []
        if validate_paths:
            value = entry.get("value")
            refs = value.get("file_paths", []) if isinstance(value, dict) else []
            if isinstance(refs, list):
                for rel in refs:
                    if isinstance(rel, str):
                        try:
                            resolved = _resolve_repo_path(rel)
                        except ValueError:
                            record["stale_paths"].append(rel)
                            continue
                        if not resolved.exists():
                            record["stale_paths"].append(rel)
        if expired or record["stale_paths"]:
            stale.append(record)
        kept.append(entry)

    if drop_expired:
        _require_mutations()
        payload["entries"] = kept
        _memory_save(payload)

    return {
        "path": str(MEMORY_FILE),
        "total_checked": len(entries),
        "stale_count": len(stale),
        "dropped_expired": dropped,
        "stale_entries": stale,
    }


async def healthz(_request):
    return JSONResponse(
        {
            "ok": True,
            "repo_path": str(REPO_PATH),
            "is_git_repo": _is_git_repo(),
            "allow_mutations": ALLOW_MUTATIONS,
            "transport": MCP_TRANSPORT,
        }
    )


async def root(_request):
    return PlainTextResponse("git-repo-manager MCP server")


@contextlib.asynccontextmanager
async def lifespan(app: Starlette):
    async with mcp.session_manager.run():
        yield


starlette_app = Starlette(
    routes=[
        Route("/", root, methods=["GET"]),
        Route("/healthz", healthz, methods=["GET"]),
        Mount("/", app=mcp.streamable_http_app()),
    ],
    lifespan=lifespan,
)

app = CORSMiddleware(
    starlette_app,
    allow_origins=ALLOW_ORIGINS,
    allow_methods=["GET", "POST", "DELETE"],
    allow_headers=["*"],
    expose_headers=["Mcp-Session-Id"],
)


def main() -> None:
    transport = MCP_TRANSPORT

    if transport in {"stdio", "direct"}:
        mcp.run()
        return

    if transport in {"http", "streamable-http", "streamable_http"}:
        import uvicorn

        uvicorn.run(app, host=HOST, port=PORT)
        return

    raise ValueError(
        "Unsupported MCP_TRANSPORT. Expected one of: stdio, direct, http, streamable-http"
    )


if __name__ == "__main__":
    main()
